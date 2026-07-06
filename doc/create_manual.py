#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成隔离泄漏试验系统操作指南 Word 文档 - 基于实际代码
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='宋体', size=None, bold=False, color=None):
    """设置字体属性"""
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_screenshot_placeholder(doc, description, page_name):
    """添加截图占位符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n【截图位置：{page_name}】\n{description}\n")
    set_font(run, size=Pt(11), color=RGBColor(255, 0, 0))
    run.font.italic = True
    p.space_after = Pt(12)

def add_paragraph_with_font(doc, text, font_name='宋体', size=Pt(11), bold=False, color=RGBColor(0, 0, 0)):
    """添加段落并设置字体"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, font_name, size, bold, color)
    return p

def add_heading_with_font(doc, text, level=1, font_name='宋体', size=None, color=RGBColor(0, 0, 0)):
    """添加标题并设置字体"""
    heading = doc.add_heading(text, level=level)
    if size is None:
        if level == 1:
            size = Pt(16)
        elif level == 2:
            size = Pt(14)
        else:
            size = Pt(12)
    for run in heading.runs:
        set_font(run, font_name, size, bold=True, color=color)
    return heading

def create_manual():
    doc = Document()

    # 设置默认字体 - 黑色宋体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置标题样式 - 黑色宋体
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '宋体'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置列表样式 - 黑色宋体
    for style_name in ['List Bullet', 'List Number', 'List Bullet 2', 'List Number 2']:
        try:
            list_style = doc.styles[style_name]
            list_style.font.name = '宋体'
            list_style.font.color.rgb = RGBColor(0, 0, 0)
            list_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except:
            pass

    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    title = add_heading_with_font(doc, '隔离阀泄漏试验管理系统', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size=Pt(26), bold=True)

    subtitle = add_heading_with_font(doc, '用户操作指南', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, size=Pt(18))

    doc.add_paragraph()
    version = doc.add_paragraph('版本：V1.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========== 目录占位 ==========
    add_heading_with_font(doc, '目录', level=1)
    p = doc.add_paragraph()
    run = p.add_run('\n【请在此处插入自动目录】\n\n操作方法：\n1. 删除本段文字\n2. 点击Word菜单：引用 → 目录 → 自动目录\n3. Word会根据标题样式自动生成目录\n')
    set_font(run, size=Pt(11), color=RGBColor(255, 0, 0))
    run.font.italic = True

    doc.add_page_break()

    # ========== 1. 系统概述 ==========
    add_heading_with_font(doc, '1. 系统概述', level=1)
    add_paragraph_with_font(doc, '本系统用于管理隔离阀泄漏试验的全流程，主要功能包括：')
    items = [
        '概览页面：系统使用统计和核心指标展示',
        '试验记录：查看、搜索、管理试验记录，查看过程曲线',
        '数据上传：单文件导入试验数据',
        '批量上传：从文件夹批量导入试验数据',
        '基础数据：项目/机组/试验对象/测量装置台账管理，报告导出',
        '实时监测：连接PLC设备，实时采集和显示试验数据',
        '系统管理：用户管理、角色权限、操作日志、数据备份',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_paragraph_with_font(doc, '系统适用于核电站隔离阀泄漏试验的数据管理场景，支持从试验装置导出的CSV/JSON数据文件批量导入，自动生成试验记录并关联过程曲线数据。')

    # ========== 2. 登录与主界面 ==========
    add_heading_with_font(doc, '2. 登录与主界面', level=1)

    add_heading_with_font(doc, '2.1 系统登录', level=2)
    add_paragraph_with_font(doc, '启动系统后，会显示登录界面：')
    add_paragraph_with_font(doc, '• 输入用户名')
    add_paragraph_with_font(doc, '• 输入密码')
    add_paragraph_with_font(doc, '• 点击"登录"按钮')

    add_screenshot_placeholder(doc, '登录界面截图', '登录页面')

    add_heading_with_font(doc, '2.2 主界面介绍', level=2)
    add_paragraph_with_font(doc, '登录成功后进入主界面，左侧为功能导航菜单，右侧为功能内容区域：')

    add_screenshot_placeholder(doc, '主界面全貌截图，显示左侧导航菜单和右侧内容区域', '主界面/概览页面')

    add_paragraph_with_font(doc, '左侧导航菜单包含：')
    menu_items = [
        ('概览', '系统使用统计和核心指标展示'),
        ('试验记录', '查看、搜索、管理试验记录，查看过程曲线'),
        ('数据上传', '单文件导入试验数据'),
        ('批量上传', '从文件夹批量导入试验数据'),
        ('基础数据', '项目/机组/试验对象/测量装置台账管理'),
        ('实时监测', '连接PLC设备，实时采集试验数据'),
        ('系统管理', '用户管理、角色权限、操作日志、数据备份'),
    ]
    for name, desc in menu_items:
        add_paragraph_with_font(doc, f'• {name}：{desc}')

    doc.add_page_break()

    # ========== 3. 概览页面 ==========
    add_heading_with_font(doc, '3. 概览页面', level=1)
    add_paragraph_with_font(doc, '点击左侧菜单"概览"进入概览页面，该页面为纯展示页面，无可操作按钮。')

    add_screenshot_placeholder(doc, '概览页面截图，显示统计卡片和列表', '概览页面')

    add_paragraph_with_font(doc, '页面包含：')
    overview_items = [
        '核心指标卡片：试验对象数、测量装置数、试验记录数、合格率、异常数、备份状态',
        '最近导入的试验记录列表',
        '最新一条导入的详细信息',
        '台账概况：项目数、机组数、系统数、贯穿件数、阀门数等统计',
        '装置状态：各装置连接状态列表',
        '系统维护：数据库连接状态、最近备份时间等',
    ]
    for item in overview_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 4. 试验记录管理 ==========
    add_heading_with_font(doc, '4. 试验记录管理', level=1)

    add_heading_with_font(doc, '4.1 查看试验记录列表', level=2)
    add_paragraph_with_font(doc, '点击左侧菜单"试验记录"进入试验记录列表页面。')

    add_screenshot_placeholder(doc, '试验记录列表页面截图，显示筛选栏、列表和底部曲线区域', '试验记录列表页面')

    add_paragraph_with_font(doc, '页面分为三部分：')
    add_paragraph_with_font(doc, '• 顶部：筛选条件栏')
    add_paragraph_with_font(doc, '• 中部：试验记录列表（DataGrid）')
    add_paragraph_with_font(doc, '• 底部：过程曲线和配方参数显示区')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '列表显示以下信息：')
    fields = [
        '序号、记录编号、项目、机组、对象编码',
        '最终泄漏率、泄漏限值、结果（合格/不合格）',
        '使用配方、测量装置、操作人员、试验时间、导入时间',
    ]
    for f in fields:
        add_paragraph_with_font(doc, f'• {f}')

    add_heading_with_font(doc, '4.2 搜索和筛选试验记录', level=2)
    add_paragraph_with_font(doc, '在列表页面上方有筛选功能：')

    add_screenshot_placeholder(doc, '筛选区域截图，显示筛选条件和按钮', '试验记录列表页面的筛选区域')

    add_paragraph_with_font(doc, '第一行筛选条件：')
    filters1 = [
        '项目：下拉选择特定项目',
        '机组：下拉选择特定机组（跟随项目联动）',
        '结果：全部 / 合格 / 不合格',
    ]
    for f in filters1:
        add_paragraph_with_font(doc, f'• {f}')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '第二行筛选条件：')
    filters2 = [
        '日期范围：选择起止日期',
        '关键字：模糊匹配记录编号、试验对象、测量装置、数据包名称',
        '查询按钮：执行筛选查询',
        '重置按钮：清空所有筛选条件',
    ]
    for f in filters2:
        add_paragraph_with_font(doc, f'• {f}')

    add_heading_with_font(doc, '4.3 查看试验详情和过程曲线', level=2)
    add_paragraph_with_font(doc, '在列表中单击选中一条记录，底部区域会自动显示：')

    add_screenshot_placeholder(doc, '选中记录后底部显示曲线和配方参数的截图', '试验记录详情（底部区域）')

    add_paragraph_with_font(doc, '• 左侧：过程曲线回放（压力、流量、温度等随时间变化的曲线图）')
    add_paragraph_with_font(doc, '• 右侧：通道图例 + 配方参数面板（配方名称、气密目标压、精吹目标压、预期泄漏流量）')

    add_heading_with_font(doc, '4.4 修改试验配方', level=2)
    add_paragraph_with_font(doc, '双击数据行，或选中记录后点击"批量修改配方"按钮：')

    add_screenshot_placeholder(doc, '配方修改对话框截图', '配方修改弹窗')

    add_paragraph_with_font(doc, '• 弹出配方修改对话框')
    add_paragraph_with_font(doc, '• 选择新的试验配方')
    add_paragraph_with_font(doc, '• 点击确认，更新该记录的配方和配方快照')

    add_heading_with_font(doc, '4.5 批量修改配方', level=2)
    add_paragraph_with_font(doc, '1. 在列表中勾选多条记录（支持跨页保持选中状态）')
    add_paragraph_with_font(doc, '2. 点击"批量修改配方"按钮')
    add_paragraph_with_font(doc, '3. 在弹窗中选择新配方')
    add_paragraph_with_font(doc, '4. 点击确认，批量更新所有选中记录的配方')

    add_heading_with_font(doc, '4.6 批量删除试验记录', level=2)
    add_paragraph_with_font(doc, '1. 在列表中勾选要删除的记录（可多选）')
    add_paragraph_with_font(doc, '2. 点击"批量删除"按钮')
    add_paragraph_with_font(doc, '3. 系统弹出确认对话框，确认后记录将被删除')
    add_paragraph_with_font(doc, '注意：此操作不可恢复！')

    add_screenshot_placeholder(doc, '删除确认对话框截图', '点击批量删除后弹出的确认框')

    doc.add_page_break()

    # ========== 5. 数据导入 ==========
    add_heading_with_font(doc, '5. 数据上传', level=1)

    add_heading_with_font(doc, '5.1 单文件上传', level=2)
    add_paragraph_with_font(doc, '点击左侧菜单"数据上传"进入单文件导入页面。')

    add_screenshot_placeholder(doc, '数据上传页面截图，显示左侧表单和右侧说明', '数据上传页面')

    add_paragraph_with_font(doc, '操作步骤：')
    steps = [
        '点击"选择数据包文件"按钮，选择数据文件（支持CSV/JSON/TXT）',
        '系统自动解析文件内容并加载配方列表',
        '填写表单：记录编号、项目编码、机组编码、操作人员',
        '选择试验配方（系统可能自动匹配试验对象的默认配方）',
        '点击"开始上传"按钮完成导入',
    ]
    for i, step in enumerate(steps, 1):
        add_paragraph_with_font(doc, f'{i}. {step}')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '页面右侧显示数据结构说明和最近上传记录列表。')

    add_heading_with_font(doc, '5.2 批量上传（文件夹）', level=2)
    add_paragraph_with_font(doc, '点击左侧菜单"批量上传"进入批量上传页面。')

    add_screenshot_placeholder(doc, '批量上传页面截图，显示文件列表和匹配状态', '批量上传页面')

    add_paragraph_with_font(doc, '操作步骤：')
    steps = [
        '点击"选择文件夹"按钮，选择数据文件夹',
        '系统自动扫描文件夹，解析所有数据文件',
        '在列表中查看解析结果，确认项目/机组/试验对象匹配正确',
        '就绪的文件显示绿色背景，待补充的文件显示黄色背景',
        '点击"开始上传"按钮',
        '等待上传完成，查看导入结果',
    ]
    for i, step in enumerate(steps, 1):
        add_paragraph_with_font(doc, f'{i}. {step}')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '文件夹结构要求：')
    add_paragraph_with_font(doc, '数据文件夹应按以下层级组织：')
    structure = [
        '根文件夹/',
        '  ├── 项目名称/',
        '  │    └── 机组名称/',
        '  │         └── 系统名称/',
        '  │              └── 贯穿件编号/',
        '  │                   └── 阀门编号/',
        '  │                        ├── xxx_结果汇总.csv',
        '  │                        └── xxx_过程数据.csv',
    ]
    for line in structure:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.5)

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, 'CSV文件格式说明：')
    add_paragraph_with_font(doc, '• 结果汇总CSV：包含试验对象编码、装置编号、试验时间、试验压力、泄漏率、判定结果等元数据')
    add_paragraph_with_font(doc, '• 过程数据CSV：包含时间序列的压力、流量、温度等通道数据')
    add_paragraph_with_font(doc, '• 系统会自动将同名的结果汇总CSV和过程数据CSV配对，创建完整的试验记录')

    doc.add_page_break()

    # ========== 6. 基础数据管理 ==========
    add_heading_with_font(doc, '6. 基础数据管理', level=1)
    add_paragraph_with_font(doc, '点击左侧菜单"基础数据"进入基础数据管理页面，包含4个标签页。')

    add_screenshot_placeholder(doc, '基础数据页面截图，显示标签页切换', '基础数据页面')

    add_heading_with_font(doc, '6.1 项目/机组管理', level=2)
    add_paragraph_with_font(doc, '管理试验所属的项目和机组信息。')

    add_screenshot_placeholder(doc, '项目/机组管理页面截图，显示左侧项目列表和右侧机组列表', '基础数据-项目/机组标签页')

    add_paragraph_with_font(doc, '左侧 - 项目管理：')
    add_paragraph_with_font(doc, '• 新增：点击"新增"按钮，填写项目编码、名称、状态')
    add_paragraph_with_font(doc, '• 编辑：选中项目后点击"编辑"，修改信息')
    add_paragraph_with_font(doc, '• 删除：选中项目后点击"删除"')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '右侧 - 机组管理（联动左侧选中的项目）：')
    add_paragraph_with_font(doc, '• 新增：点击"新增"按钮，填写机组编码、名称、状态')
    add_paragraph_with_font(doc, '• 编辑：选中机组后点击"编辑"，修改信息')
    add_paragraph_with_font(doc, '• 删除：选中机组后点击"删除"')

    add_heading_with_font(doc, '6.2 试验对象管理', level=2)
    add_paragraph_with_font(doc, '管理试验对象的层级结构：系统 → 贯穿件 → 阀门/部件。')

    add_screenshot_placeholder(doc, '试验对象管理页面截图，显示左侧路径树和右侧详情', '基础数据-试验对象标签页')

    add_paragraph_with_font(doc, '左侧 - 路径树：')
    add_paragraph_with_font(doc, '• 树形结构展示系统、贯穿件、阀门的层级关系')
    add_paragraph_with_font(doc, '• 底部快速新建按钮：新建系统、新建贯穿件、新建阀门、新建其他部件')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '右侧 - 节点详情：')
    add_paragraph_with_font(doc, '• 显示选中节点的详细信息（编号、名称、类型、泄漏率限值、试验压力等）')
    add_paragraph_with_font(doc, '• 操作按钮：修改节点、导入数据、导出数据、删除节点')
    add_paragraph_with_font(doc, '• 下方显示统计概览（累计试验次数、合格/不合格次数）')
    add_paragraph_with_font(doc, '• 下方显示关联配方信息')

    add_heading_with_font(doc, '6.3 测量装置管理', level=2)
    add_paragraph_with_font(doc, '管理试验使用的测量装置。')

    add_screenshot_placeholder(doc, '测量装置管理页面截图', '基础数据-测量装置标签页')

    add_paragraph_with_font(doc, '操作：')
    add_paragraph_with_font(doc, '• 新增装置：点击"新增装置"按钮')
    add_paragraph_with_font(doc, '• 编辑装置：选中装置后点击"编辑装置"')
    add_paragraph_with_font(doc, '• 删除装置：选中装置后点击"删除"')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '装置信息包括：装置编号、装置名称、型号、序列号、主通信方式、启用状态等。')

    add_heading_with_font(doc, '6.4 报告导出', level=2)
    add_paragraph_with_font(doc, '导出试验数据和报告。')

    add_screenshot_placeholder(doc, '报告导出页面截图', '基础数据-报告导出标签页')

    add_paragraph_with_font(doc, '操作步骤：')
    steps = [
        '选择导出范围',
        '选择导出格式',
        '输入文件名（留空自动生成）',
        '选择导出目录',
        '点击"导出报告"按钮',
    ]
    for i, step in enumerate(steps, 1):
        add_paragraph_with_font(doc, f'{i}. {step}')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '快速导出按钮：')
    add_paragraph_with_font(doc, '• 导出Excel：快速导出为Excel格式')
    add_paragraph_with_font(doc, '• 导出PDF：快速导出为PDF格式')

    doc.add_page_break()

    # ========== 7. 实时监测 ==========
    add_heading_with_font(doc, '7. 实时监测', level=1)

    add_paragraph_with_font(doc, '点击左侧菜单"实时监测"进入实时监测页面。')

    add_screenshot_placeholder(doc, '实时监测页面截图，显示控制区、变量表格和趋势曲线', '实时监测页面')

    add_heading_with_font(doc, '7.1 配置监测范围', level=2)
    add_paragraph_with_font(doc, '顶部第一行 - 范围选择：')
    add_paragraph_with_font(doc, '• 项目：选择项目')
    add_paragraph_with_font(doc, '• 机组：选择机组')
    add_paragraph_with_font(doc, '• 试验对象：选择要监测的对象')
    add_paragraph_with_font(doc, '• 配方：选择关联的试验配方')

    add_heading_with_font(doc, '7.2 连接PLC设备', level=2)
    add_paragraph_with_font(doc, '顶部第二行 - PLC连接控制：')
    add_paragraph_with_font(doc, '• PLC地址：输入PLC的IP地址')
    add_paragraph_with_font(doc, '• 保存地址：保存PLC地址配置')
    add_paragraph_with_font(doc, '• 连接PLC：建立与PLC的连接')
    add_paragraph_with_font(doc, '• 断开PLC：断开与PLC的连接')
    add_paragraph_with_font(doc, '• 连接状态：显示当前连接状态')

    add_heading_with_font(doc, '7.3 监视控制', level=2)
    add_paragraph_with_font(doc, '• 开始监视：开始实时采集数据')
    add_paragraph_with_font(doc, '• 停止监视：停止数据采集')
    add_paragraph_with_font(doc, '• 导出CSV：将当前采集的数据导出为CSV文件')

    add_heading_with_font(doc, '7.4 配置监视变量', level=2)
    add_paragraph_with_font(doc, '中部 - 实时变量表格（可编辑）：')

    add_paragraph_with_font(doc, '工具栏按钮：')
    add_paragraph_with_font(doc, '• 添加变量：添加新的监视变量')
    add_paragraph_with_font(doc, '• 删除：删除选中的变量')
    add_paragraph_with_font(doc, '• 保存配置：保存变量配置')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '可编辑列：变量名称、西门子地址、寄存器地址、数据类型、单位、最小值、最大值')
    add_paragraph_with_font(doc, '只读列：当前值、更新时间、状态')

    add_heading_with_font(doc, '7.5 查看趋势曲线', level=2)
    add_paragraph_with_font(doc, '底部 - 趋势曲线区域：')
    add_paragraph_with_font(doc, '• 左侧：实时趋势图，显示多通道曲线')
    add_paragraph_with_font(doc, '• 右侧：通道图例，显示每个变量的当前值')

    doc.add_page_break()

    # ========== 8. 系统管理 ==========
    add_heading_with_font(doc, '8. 系统管理', level=1)
    add_paragraph_with_font(doc, '点击左侧菜单"系统管理"进入系统管理页面，包含4个标签页。')

    add_screenshot_placeholder(doc, '系统管理页面截图，显示标签页切换', '系统管理页面')

    add_heading_with_font(doc, '8.1 用户管理', level=2)
    add_paragraph_with_font(doc, '管理系统用户账号。')

    add_screenshot_placeholder(doc, '用户管理页面截图，显示左侧用户列表和右侧编辑面板', '系统管理-用户管理标签页')

    add_paragraph_with_font(doc, '左侧 - 用户列表：')
    add_paragraph_with_font(doc, '• 搜索框：实时过滤用户')
    add_paragraph_with_font(doc, '• 新增按钮：添加新用户')
    add_paragraph_with_font(doc, '• 刷新按钮：刷新列表')
    add_paragraph_with_font(doc, '• 列表显示：用户名、昵称、最后登录、状态、备注')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '右侧 - 用户编辑面板：')
    add_paragraph_with_font(doc, '• 选中用户后显示编辑表单')
    add_paragraph_with_font(doc, '• 可编辑：用户名、昵称、密码（留空不改）、状态')
    add_paragraph_with_font(doc, '• 角色分配：多选CheckBox列表，按角色勾选')
    add_paragraph_with_font(doc, '• 保存/取消按钮')

    add_heading_with_font(doc, '8.2 角色管理', level=2)
    add_paragraph_with_font(doc, '管理系统角色及其权限。')

    add_screenshot_placeholder(doc, '角色管理页面截图', '系统管理-角色管理标签页')

    add_paragraph_with_font(doc, '预置角色：')
    roles = [
        ('admin', '拥有所有权限'),
        ('operator', '可进行数据导入、查看、导出等操作'),
        ('viewer', '只能查看数据，不能修改'),
    ]
    for name, desc in roles:
        add_paragraph_with_font(doc, f'• {name}：{desc}')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '可新增自定义角色并配置权限。')

    add_heading_with_font(doc, '8.3 操作日志', level=2)
    add_paragraph_with_font(doc, '查看系统操作日志，记录所有用户的操作行为。')

    add_screenshot_placeholder(doc, '操作日志页面截图', '系统管理-操作日志标签页')

    add_paragraph_with_font(doc, '顶部 - 过滤条件：')
    add_paragraph_with_font(doc, '• 操作类型：选择特定操作类型')
    add_paragraph_with_font(doc, '• 搜索框：关键字搜索')
    add_paragraph_with_font(doc, '• 日期范围：选择起止日期')
    add_paragraph_with_font(doc, '• 查询/重置按钮')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '中部 - 日志列表：')
    add_paragraph_with_font(doc, '• 显示：操作类型、用户名、IP地址、操作时间、结果、操作详情')
    add_paragraph_with_font(doc, '• 双击日志行可查看详情')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '底部 - 日志清理：')
    add_paragraph_with_font(doc, '• 保留天数：设置日志保留天数（默认90天）')
    add_paragraph_with_font(doc, '• 预览清理：查看可清理的日志条数')
    add_paragraph_with_font(doc, '• 执行清理：清理过期日志')
    add_paragraph_with_font(doc, '• 导出当前范围：导出日志为CSV文件')

    add_heading_with_font(doc, '8.4 数据备份', level=2)
    add_paragraph_with_font(doc, '管理数据库备份。')

    add_screenshot_placeholder(doc, '数据备份页面截图', '系统管理-数据备份标签页')

    add_paragraph_with_font(doc, '顶部操作区：')
    add_paragraph_with_font(doc, '• 立即备份：执行一次数据库备份')
    add_paragraph_with_font(doc, '• 还原数据库：从备份文件还原数据库')
    add_paragraph_with_font(doc, '• 状态卡片：显示最后备份时间、数据库大小')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '备份配置区：')
    add_paragraph_with_font(doc, '• 默认备份路径')
    add_paragraph_with_font(doc, '• 备份保留策略')
    add_paragraph_with_font(doc, '• 自动备份：启用/禁用自动备份，设置间隔时间')

    add_paragraph_with_font(doc, '')
    add_paragraph_with_font(doc, '备份历史列表：')
    add_paragraph_with_font(doc, '• 显示：文件名、大小、创建时间、最后修改时间')

    # 保存文档
    output_path = r'F:\workspace\cechuang\projects\Isolation\doc\隔离阀泄漏试验系统操作指南.docx'
    doc.save(output_path)
    print(f'操作指南已生成：{output_path}')

if __name__ == '__main__':
    create_manual()
