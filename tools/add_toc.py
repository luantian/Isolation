# -*- coding: utf-8 -*-
"""给用户操作手册添加目录"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document('doc/用户操作手册.docx')

# 找到第一个 H2 的位置（目录插入点）
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2':
        insert_idx = i
        break
print(f'目录插入位置：段落 [{insert_idx}] 前')

# 收集所有标题条目
toc_entries = []
for p in doc.paragraphs:
    if p.style.name in ['Heading 1', 'Heading 2', 'Heading 3']:
        level = int(p.style.name.split()[-1])
        toc_entries.append((level, p.text.strip()))
print(f'标题条目：{len(toc_entries)} 个')

# 在文档末尾构建目录内容
body = doc.element.body

# 目录标题
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_title.add_run('目  录')
run.font.name = '宋体'
run.font.size = Pt(18)
run.bold = True

# 空行
doc.add_paragraph()

# TOC 域代码（Word 打开后按 Ctrl+A 再 F9 可更新）
toc_field = doc.add_paragraph()
# begin
r1 = toc_field.add_run()
fc_begin = r1._element.makeelement(qn('w:fldChar'), {})
fc_begin.set(qn('w:fldCharType'), 'begin')
r1._element.append(fc_begin)
# instrText
r2 = toc_field.add_run()
instr = r2._element.makeelement(qn('w:instrText'), {})
instr.set(qn('xml:space'), 'preserve')
instr.text = r' TOC \o "1-3" \h \z \u '
r2._element.append(instr)
# separate
r3 = toc_field.add_run()
fc_sep = r3._element.makeelement(qn('w:fldChar'), {})
fc_sep.set(qn('w:fldCharType'), 'separate')
r3._element.append(fc_sep)
# end
r4 = toc_field.add_run()
fc_end = r4._element.makeelement(qn('w:fldChar'), {})
fc_end.set(qn('w:fldCharType'), 'end')
r4._element.append(fc_end)

# 空行
doc.add_paragraph()

# 手动目录条目（作为备用，域更新后会被替换）
for level, text in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20 * (level - 1))
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(14) if level == 1 else Pt(12)
    if level == 1:
        run.bold = True
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = r2._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rPr)
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:hAnsi'), '宋体')
    rPr.insert(0, rFonts)

# 空行 + 分页符
doc.add_paragraph()
page_break = doc.add_paragraph()
run_pb = page_break.add_run()
br = run_pb._element.makeelement(qn('w:br'), {})
br.set(qn('w:type'), 'page')
run_pb._element.append(br)

# 收集刚添加的所有元素（从 toc_title 开始）
toc_elements = []
found_toc = False
for elem in list(body):
    if elem is toc_title._element:
        found_toc = True
    if found_toc:
        toc_elements.append(elem)

print(f'目录元素：{len(toc_elements)} 个')

# 移动到 insert_idx 位置前
target = doc.paragraphs[insert_idx]._element
for elem in toc_elements:
    body.remove(elem)
    target.addprevious(elem)

print(f'已移动目录到段落 [{insert_idx}] 前')

doc.save('doc/用户操作手册_toc.docx')
print('已保存到：doc/用户操作手册_toc.docx')

# 验证
doc2 = Document('doc/用户操作手册.docx')
print('\n=== 验证 ===')
for i in range(min(30, len(doc2.paragraphs))):
    p = doc2.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f'[{i}] [{p.style.name}] {t[:50]}')
