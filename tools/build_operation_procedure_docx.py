from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = "doc/智能安全壳隔离阀泄漏率数据管理软件操作规程.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4)
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, after=0, line=1.15)
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            for p in cells[i].paragraphs:
                set_paragraph_spacing(p, after=0, line=1.15)
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("智能安全壳隔离阀泄漏率数据管理软件\n操作规程")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("适用于数据库管理软件的日常操作、数据维护、上传下载、查询分析和备份管理")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")

    rows = [
        ("文档编号", "待定"),
        ("版本", "V1.0"),
        ("适用范围", "智能安全壳隔离阀泄漏率数据管理软件"),
        ("编制日期", "2026年5月"),
        ("使用对象", "系统管理员、数据管理员、试验工程师、数据分析人员"),
    ]
    add_table(doc, ["项目", "内容"], rows, [2700, 6660])
    doc.add_page_break()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "智能安全壳隔离阀泄漏率数据管理软件操作规程"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer._p.append(fld_begin)
    footer._p.append(instr)
    footer._p.append(fld_end)
    footer.add_run(" 页")

    add_title_page(doc)

    doc.add_heading("1. 目的", level=1)
    add_paragraph(doc, "为规范智能安全壳隔离阀泄漏率数据管理软件的使用过程，保证试验对象台账、测量装置数据、试验记录、统计分析结果和数据库备份等信息准确、完整、可追溯，特制定本操作规程。")

    doc.add_heading("2. 适用范围", level=1)
    add_paragraph(doc, "本规程适用于数据管理软件的日常操作，包括用户登录、基础台账维护、测量装置管理、试验任务下发、试验数据上传、数据查询分析、报告导出、数据备份恢复和系统日志查看。")
    add_paragraph(doc, "本规程不适用于测量装置硬件操作、传感器校准、PID 过程控制、泄漏率计算模型验证和试验台架操作。上述内容应按对应设备或系统文件执行。")

    doc.add_heading("3. 术语和缩略语", level=1)
    add_table(doc, ["术语", "说明"], [
        ("试验对象", "需要进行泄漏率试验的阀门、贯穿件或其他密封性部件。"),
        ("试验路径", "按项目、机组、系统、贯穿件、阀门或部件形成的数据层级路径。"),
        ("测量装置", "智能安全壳隔离阀泄漏率测量装置，用于现场采集试验数据。"),
        ("任务下发", "将软件中配置的试验对象和试验参数下载到测量装置。"),
        ("数据上传", "将测量装置产生的试验结果和过程数据上传至数据库。"),
        ("历史记录", "同一试验对象历次试验形成的记录集合，上传新数据不得覆盖既有记录。"),
    ], [2200, 7160])

    doc.add_heading("4. 角色与职责", level=1)
    add_table(doc, ["角色", "职责"], [
        ("系统管理员", "负责账号、角色、权限、备份恢复、数据库迁移、系统日志和基础参数维护。"),
        ("数据管理员", "负责项目、机组、系统、贯穿件、阀门、部件及测量装置台账维护。"),
        ("试验工程师", "负责选择试验对象、配置任务、下发任务、上传试验数据并核对结果。"),
        ("数据分析人员", "负责查询历史记录、查看趋势图、统计合格情况并导出分析结果。"),
        ("只读用户", "负责查看授权范围内的数据和报告，不进行新增、修改、删除和维护操作。"),
    ], [2100, 7260])

    doc.add_heading("5. 操作前准备", level=1)
    add_bullets(doc, [
        "确认计算机、数据库服务和数据管理软件运行正常。",
        "确认当前用户已获得相应操作权限。",
        "确认项目、机组、系统、贯穿件、阀门或部件基础台账已建立，或具备维护权限。",
        "进行任务下发或数据上传前，确认测量装置已按约定方式连接，并处于可通信状态。",
        "进行数据库恢复、迁移等高权限操作前，应先完成现有数据备份，并取得授权。",
    ])

    doc.add_heading("6. 总体操作流程", level=1)
    add_numbers(doc, [
        "用户登录软件，确认当前角色和权限。",
        "维护或核对项目、机组、系统、贯穿件、阀门及其他部件台账。",
        "登记或核对测量装置信息。",
        "选择试验对象路径，生成试验任务。",
        "将试验任务下载至测量装置。",
        "现场试验完成后，从测量装置上传试验数据。",
        "核对上传结果，确认数据已按时间顺序保存且未覆盖历史记录。",
        "查询试验记录，查看泄漏率趋势、合格情况和统计分析结果。",
        "按需要导出报告或数据文件。",
        "按备份策略执行数据库备份，并定期检查备份可用性。",
    ])

    doc.add_heading("7. 账号登录与退出", level=1)
    doc.add_heading("7.1 登录", level=2)
    add_numbers(doc, [
        "启动数据管理软件。",
        "在登录界面输入用户名和密码。",
        "点击“登录”。",
        "登录成功后，确认首页显示的用户名称、角色和系统状态信息正确。",
    ])
    doc.add_heading("7.2 退出", level=2)
    add_numbers(doc, [
        "确认当前数据维护、上传、下载或导出操作已完成。",
        "点击“退出登录”或关闭软件。",
        "如系统提示存在未完成操作，应先完成或取消相关操作后再退出。",
    ])

    doc.add_heading("8. 基础台账维护", level=1)
    doc.add_heading("8.1 项目和机组维护", level=2)
    add_numbers(doc, [
        "进入“资产管理”或“基础台账”模块。",
        "新建项目，填写项目名称、项目编号、备注等信息。",
        "在项目下新建机组，填写机组号、机组名称和备注。",
        "保存后检查项目和机组是否出现在左侧路径树或列表中。",
    ])
    doc.add_heading("8.2 系统、贯穿件和部件路径维护", level=2)
    add_numbers(doc, [
        "选择目标项目和机组。",
        "新建工艺系统，填写系统名称和系统编码。",
        "根据现场管理需要，在系统下新建贯穿件路径，也可直接新建阀门或其他密封性部件路径。",
        "保存后检查层级关系是否正确。",
    ])
    doc.add_heading("8.3 阀门和其他部件维护", level=2)
    add_numbers(doc, [
        "选择对应路径，点击“新建阀门”或“新建部件”。",
        "填写设备位号、部件类型、阀门类型、泄漏率限值、试验压力等信息。",
        "核对关键参数，确认无误后保存。",
        "如需修改，选择目标对象后点击“编辑”；如需删除，应确认该对象无有效历史数据或按客户管理要求执行。",
    ])

    doc.add_heading("9. 测量装置管理", level=1)
    add_numbers(doc, [
        "进入“测量装置管理”模块。",
        "点击“新增装置”，填写装置编号、名称、型号、序列号、通信方式和备注。",
        "保存后检查装置是否进入装置列表。",
        "连接测量装置后，查看状态是否显示为在线或最近同步时间是否更新。",
        "如装置更换或停用，应及时更新装置状态，保留历史数据来源信息。",
    ])

    doc.add_heading("10. 试验任务下发", level=1)
    add_numbers(doc, [
        "进入“任务管理”或“任务下发”模块。",
        "选择需要下发的试验对象路径，可按系统、贯穿件、单个阀门或其他部件选择。",
        "检查所选对象的泄漏率限值、试验压力、阀门类型等关键参数是否完整。",
        "选择目标测量装置。",
        "点击“生成任务”并确认任务内容。",
        "点击“下载至装置”。",
        "下载完成后查看系统提示，确认任务下发成功。",
        "如下载失败，应查看失败原因，检查装置连接、权限、任务参数和通信配置后重试。",
    ])

    doc.add_heading("11. 试验数据上传", level=1)
    add_numbers(doc, [
        "确认测量装置已完成现场试验并保存试验数据。",
        "连接测量装置，进入“数据上传”模块。",
        "选择需要上传的数据范围，可按试验对象路径或测量装置数据包选择。",
        "点击“读取数据”或“上传数据”。",
        "系统解析数据后，核对试验对象、测量装置编号、试验时间、最终泄漏率和判定结果。",
        "确认无误后执行入库。",
        "上传完成后，进入试验记录列表，检查新增记录是否按时间顺序保存。",
        "同一试验对象再次上传时，系统应新增历史记录，不得覆盖既有记录。",
    ])

    doc.add_heading("12. 试验记录查询", level=1)
    add_numbers(doc, [
        "进入“数据中心”或“试验记录查询”模块。",
        "按项目、机组、系统、贯穿件、阀门、部件、测量装置、试验时间、结果状态等条件筛选。",
        "点击“查询”。",
        "在列表中查看试验时间、最终泄漏率、判定结果、操作人员和测量装置来源。",
        "点击单条记录查看详情，包括基础信息、过程数据和上传日志。",
    ])

    doc.add_heading("13. 数据分析与趋势查看", level=1)
    add_numbers(doc, [
        "进入“分析中心”模块。",
        "选择分析对象，可选择单个阀门、部件、系统、机组或阀门类型。",
        "设置时间范围和筛选条件。",
        "查看单个阀门泄漏率历史趋势图。",
        "查看不同类型阀门泄漏率曲线对比。",
        "查看故障类型趋势、试验次数、单次合格率和整台机组阀门合格情况。",
        "对不合格记录进行定位，查看对应试验详情和过程数据。",
    ])

    doc.add_heading("14. 报告与数据导出", level=1)
    add_numbers(doc, [
        "在试验记录、分析结果或报告模块中选择需要导出的数据范围。",
        "选择导出格式，如 Excel、PDF 或系统支持的其他格式。",
        "确认导出内容包含项目、机组、系统、试验对象、试验时间、泄漏率、判定结果和操作人员等必要信息。",
        "点击“导出”。",
        "导出完成后检查文件内容是否完整。",
        "导出的报告和数据文件应按客户文件管理要求归档。",
    ])

    doc.add_heading("15. 数据备份、恢复与迁移", level=1)
    doc.add_heading("15.1 手动备份", level=2)
    add_numbers(doc, [
        "以系统管理员账号登录。",
        "进入“系统设置”或“数据维护”模块。",
        "选择“手动备份”。",
        "确认备份路径和备份名称。",
        "执行备份并等待系统提示备份成功。",
        "记录备份时间、操作人员和备份文件位置。",
    ])
    doc.add_heading("15.2 定期备份", level=2)
    add_numbers(doc, [
        "进入备份策略配置页面。",
        "设置备份周期、执行时间、保存位置和保留期限。",
        "保存配置。",
        "定期检查备份任务执行日志，确认备份文件生成正常。",
    ])
    doc.add_heading("15.3 数据恢复", level=2)
    add_numbers(doc, [
        "确认恢复操作已获得授权。",
        "恢复前先备份当前数据库。",
        "选择目标备份文件。",
        "执行恢复操作。",
        "恢复完成后，抽查项目台账、试验记录、过程数据、用户权限和日志是否完整。",
    ])
    doc.add_heading("15.4 数据迁移", level=2)
    add_numbers(doc, [
        "确认源数据库和目标数据库环境可用。",
        "执行源数据库备份。",
        "在目标环境导入备份数据或执行迁移工具。",
        "迁移完成后核对基础台账、历史试验记录、过程数据、用户权限和日志。",
        "完成迁移记录，注明迁移时间、操作人员、源环境和目标环境。",
    ])

    doc.add_heading("16. 异常处理", level=1)
    add_table(doc, ["异常现象", "处理要求"], [
        ("无法登录", "检查用户名、密码、账号状态和网络连接；仍无法登录时联系系统管理员。"),
        ("无操作权限", "确认当前账号角色；确需操作时按客户权限管理流程申请授权。"),
        ("装置连接失败", "检查通信线缆、网络、串口参数或 USB 连接；确认装置处于可通信状态。"),
        ("任务下载失败", "检查试验对象参数是否完整、装置是否在线、通信配置是否正确。"),
        ("数据上传失败", "检查数据包格式、试验对象是否存在、数据库连接是否正常；失败数据不得手工改写入库。"),
        ("发现重复上传", "按系统提示处理；原则上不得覆盖历史记录，应保留可追溯处理结果。"),
        ("查询结果异常", "检查筛选条件、数据权限和时间范围；必要时由管理员检查数据库状态。"),
        ("备份失败", "检查备份路径、磁盘空间、数据库连接和管理员权限，排除后重新备份。"),
    ], [2400, 6960])

    doc.add_heading("17. 安全与注意事项", level=1)
    add_bullets(doc, [
        "账号应专人专用，不得多人共用同一账号。",
        "普通用户不得执行数据库恢复、迁移、用户管理等高权限操作。",
        "试验历史记录不得随意删除或覆盖。",
        "导出文件应按客户数据管理要求保存和传递。",
        "数据库恢复、迁移和批量删除等操作应先备份，并保留操作记录。",
        "发现数据异常时，应先保留现场数据和日志，再进行问题排查。",
    ])

    doc.add_heading("18. 记录与归档", level=1)
    add_table(doc, ["记录名称", "形成环节", "保存要求"], [
        ("用户操作日志", "登录、维护、上传、下载、导出、备份等操作", "系统自动保存，按客户要求留存。"),
        ("任务下发记录", "试验任务下载至测量装置", "记录任务范围、装置编号、操作人员和时间。"),
        ("数据上传记录", "测量装置数据上传入库", "记录上传范围、上传结果、操作人员和时间。"),
        ("试验数据报告", "查询、统计或报告导出", "按项目或机组归档。"),
        ("备份恢复记录", "数据库备份、恢复、迁移", "由系统管理员保存并定期核查。"),
    ], [2100, 3660, 3600])

    doc.add_heading("附录 A 操作检查表", level=1)
    add_table(doc, ["序号", "检查项", "确认"], [
        ("1", "当前账号权限满足本次操作要求", "□"),
        ("2", "项目、机组、系统、贯穿件、阀门或部件台账已核对", "□"),
        ("3", "测量装置编号和连接状态已确认", "□"),
        ("4", "任务下发前关键试验参数已检查", "□"),
        ("5", "数据上传后确认未覆盖历史记录", "□"),
        ("6", "不合格或异常数据已完成标识和追溯", "□"),
        ("7", "报告或数据文件已按要求导出和归档", "□"),
        ("8", "数据库备份或维护操作已记录", "□"),
    ], [900, 6900, 1560])

    doc.add_heading("附录 B 版本记录", level=1)
    add_table(doc, ["版本", "日期", "修订内容", "修订人"], [
        ("V1.0", "2026年5月", "初版，形成数据管理软件操作规程。", ""),
    ], [1200, 1800, 4800, 1560])

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
