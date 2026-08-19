# -*- coding: utf-8 -*-
"""
一次性完成设计方案文档更新，从原始 A(1) 出发。
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy

INPUT = 'doc/数据分析-管理软件设计方案-A(1).docx'
OUTPUT = 'doc/数据分析-管理软件设计方案-A(2).docx'

doc = Document(INPUT)

def norm(t): return t.replace('\xa0', ' ').strip()

# ============== 1. 按索引修改段落文本 ==============
# 先建立 [index -> 新文本] 映射（精确到段落序号）
# 索引基于原始 A(1) 文档

text_by_index = {}

for i, p in enumerate(doc.paragraphs):
    t = norm(p.text)

    if i == 43 and '自动上传' in t:
        text_by_index[i] = '实现试验数据导入（单文件与批量）、自动归档、自动分类存储；'
    elif i == 51 and 'RS485' in t and 'Siemens' not in t:
        text_by_index[i] = '通信层负责实现RS232、RS485、TCP/IP、Modbus TCP、Siemens S7等协议通信；'
    elif i == 53 and 'Web管理界面' in t:
        text_by_index[i] = '应用层负责向用户提供桌面管理界面、数据分析界面、报表界面以及系统配置界面。'
    elif i == 58 and '主控台' in t:
        text_by_index[i] = '首页概览: 系统状态、核心指标及告警信息概览。'
    elif i == 59 and '资产管理中心' in t:
        text_by_index[i] = '基础台账（试验对象）: 用于管理项目、机组、试验对象路径树、测量装置等物理资产的数字模型，以及试验报告导出。'
    elif i == 60 and '任务管理中心' in t:
        text_by_index[i] = '试验路径（配方管理）: 用于创建和管理试验路径配方，配置泄漏率限值、阀门参数等试验条件。'
    elif i == 61 and '数据中心' in t:
        text_by_index[i] = '试验记录: 用于存储、查询和管理所有历史测试记录，支持曲线回放与批量操作。'
    elif i == 62 and '分析中心' in t:
        text_by_index[i] = '实时监视: 连接PLC实时采集数据，以趋势曲线形式展示压力、温度、流量等参数变化。'
    elif i == 63 and '系统设置' in t and '用户管理' in t:
        text_by_index[i] = '数据分析: 提供故障趋势、合格率统计、泄漏率趋势等多维度数据分析工具与图表。'
    elif i == 70 and '装置注册与管理' in t:
        text_by_index[i] = '装置注册与管理：支持添加、删除、编辑多台测量装置信息（编号、IP、通信方式、启用状态等）。试验路径配方管理：支持配方的创建、编辑、删除，每次编辑自动生成版本快照，支持配方CSV导入导出。'
    elif i == 72 and '贯穿件/部件' in t:
        text_by_index[i] = '机组 (如：海南3号机组、漳州3号机组) → 系统 (如：RHR、CPS) → 贯穿件 (如：PN101) → 阀门/其他部件 (如：1RHR040VP)'
    elif i == 74 and t == '贯穿件泄漏率限值':
        text_by_index[i] = '节点泄漏率限值（叶子节点用于合格判定）'
    elif i == 75 and '阀门泄漏率' in t:
        text_by_index[i] = '试验压力、阀门类型（电动阀、止回阀、闸阀等）'
    elif i == 76 and '其他部件' in t:
        text_by_index[i] = '默认关联试验路径（叶子节点可指定默认配方）'
    elif i == 78 and '双向数据交互' in t:
        text_by_index[i] = '该模块主要实现数据管理软件与多台智能安全壳隔离阀泄漏率测量装置之间的双向数据交互，支持单文件导入与批量导入，主要功能包含：'
    elif i == 79 and '试验对象下载' in t:
        text_by_index[i] = '试验任务下发：软件可将选定的试验对象（含配置参数）下发至指定的测量装置，支持选择目标装置并查看下发历史。'
    elif i == 80 and '试验数据上传' in t:
        text_by_index[i] = '试验数据上传：支持单文件导入（.json/.txt/.csv格式数据包）和批量导入（选择文件夹自动解析匹配），系统自动识别试验对象并关联试验路径配方。'
    elif i == 81 and '非覆盖存储' in t:
        text_by_index[i] = '非覆盖存储：同一试验对象的不同次试验数据将按时间顺序追加存储，确保历史数据完整。每次导入时自动保存配方快照（JSON），后续修改配方不影响历史记录。'
    elif i == 83 and '主要功能如下' in t:
        text_by_index[i] = '本模块旨在将原始试验数据转化为直观的统计图表与趋势分析结果，为阀门泄漏率评估和故障诊断提供数据支撑。系统提供五个分析维度：'
    elif i == 84 and t == '单阀门分析：':
        text_by_index[i] = '故障趋势分析：'
    elif i == 85 and '历史泄漏率' in t:
        text_by_index[i] = '按阀门类型统计合格数与不合格数，以堆叠柱状图展示故障分布。'
    elif i == 86 and '试验合格/失败' in t:
        text_by_index[i] = '合格率统计：展示总试验数、总体合格率、不合格数，并计算各阀门的合格率（≥95%为合格，80%-95%为注意，<80%为不合格）。'
    elif i == 87 and t == '机组级分析：':
        text_by_index[i] = '泄漏率趋势分析：'
    elif i == 88 and '按阀门类型' in t and '统计故障' in t:
        text_by_index[i] = '以多系列趋势曲线展示不同阀门类型的泄漏率历史变化，支持最多500个数据点。'
    elif i == 89 and '合格率仪表盘' in t:
        text_by_index[i] = '阀门试验次数统计：排名展示各阀门的试验次数（Top 50），并以柱状图展示Top 20。'
    elif i == 90 and '多因素' in t:
        text_by_index[i] = '机组合格情况统计：按机组展示总试验数、合格数、合格率，并以柱状图对比各机组合格率。所有分析结果支持导出为多Sheet的Excel文件。'
    elif i == 94 and '验证码' in t:
        text_by_index[i] = '认证：用户名+密码登录，密码加密存储，支持会话管理。'
    elif i == 95 and '审计' in t:
        text_by_index[i] = '审计：所有关键操作（登录、修改、删除、导出）记录操作日志，支持按时间范围和日志级别筛选查询，支持日志清理与导出。'
    elif i == 96 and '数据隔离' in t:
        text_by_index[i] = '数据隔离：通过角色权限控制不同用户的可见页面和可执行操作，导航栏按权限自动过滤。'
    elif i == 98 and '双服务器' in t:
        text_by_index[i] = '本模块通过自动/手动备份、主从数据库自动故障切换及灵活的数据恢复与迁移机制，确保试验数据的安全可靠与业务连续性。切换期间通过磁盘缓冲机制确保数据零丢失。'
    elif i == 99 and '每日/每周' in t:
        text_by_index[i] = '自动备份：支持可配置小时间隔定时全量备份数据库至指定存储位置，可配置备份保留天数（超期自动清理）。'
    elif i == 101 and '实时同步' in t:
        text_by_index[i] = '数据同步：主备服务器间实现数据库自动故障切换，定时健康检测（默认15秒间隔），主服务器故障时自动切换至备用服务器，主服务器恢复后自动切回。切换期间通过磁盘缓冲机制确保数据零丢失。'
    elif i == 113 and '二进制日志' in t:
        text_by_index[i] = '文件操作：实现本地试验数据文件的导入（如JSON、CSV、TXT数据包）、导出（分析报告、图表图片、原始数据包、CSV数据），以及配置文件、路径模板的读取与保存。所有文件操作均提供进度反馈与异常处理机制。'
    elif i == 114 and 'HTTP/HTTPS' in t:
        text_by_index[i] = '数据对接：封装与底层智能安全壳隔离阀泄漏率测量装置的通信协议（Modbus TCP、Siemens S7、RS232、RS485、USB），负责实时采集PLC数据、下发试验任务、接收装置上传的试验数据，并提供数据缓存与断点续传支持。同时，对外预留与第三方系统的数据交换接口。'
    elif i == 127 and '主要用于：' in t and '首页' in doc.paragraphs[i-1].text:
        text_by_index[i] = '软件首页整体采用工业监控平台风格设计，以核心指标卡片、数据表格、台账概况和系统维护四大模块为主体，主要用于：'
    elif i == 128 and t == '泄漏率试验数据管理':
        text_by_index[i] = '泄漏率试验数据管理总览'
    elif i == 129 and t == '多设备在线监控':
        text_by_index[i] = '多设备状态监控'
    elif i == 130 and t == '阀门试验结果分析':
        text_by_index[i] = '试验合格率统计'
    elif i == 133 and '所有设备的运行状态' in t:
        text_by_index[i] = '此区域主要负责全局数据概览，在程序首页即可查看6个核心指标（试验对象数量、测量装置数量、历史记录条数、本月合格率、待处理异常数、最近备份时间），以及最近导入的试验记录预览和台账概况。'
    elif i == 134 and '左侧导航 + 中部业务监控' in t:
        text_by_index[i] = '系统首页采用"左侧导航 + 顶部核心指标 + 中部业务监控 + 底部台账概况"的布局结构。该布局方式能够同时满足：数据总览、多设备状态监控、快速业务入口、试验结果分析等工业监控需求。'
    elif i == 136 and t == '泄漏率试验全过程数字化管理；':
        text_by_index[i] = '泄漏率试验关键指标一屏总览；'
    elif i == 137 and t == '多设备集中运行监控；':
        text_by_index[i] = '试验记录预览与最近导入详情；'
    elif i == 138 and t == '自动化试验数据采集；':
        text_by_index[i] = '台账概况（项目/机组/系统/贯穿件/阀门/部件数量统计）；'
    elif i == 139 and t == '历史数据统一管理；':
        text_by_index[i] = '系统维护状态（数据库连接、备份状态、同步状态）；'
    elif i == 140 and t == '设备运行状态实时监控；':
        text_by_index[i] = '核心KPI卡片（试验对象数、装置数、记录数、合格率、异常数、备份状态）；'
    elif i == 141 and t == '数据分析与趋势研究。':
        text_by_index[i] = '数据分析与趋势研究快速入口。'
    elif i == 146 and '基础台账页面用于' in t:
        text_by_index[i] = '基础台账页面用于实现数据分析-管理软件中的基础数据管理功能，是系统进行试验对象管理、设备管理及数据分类管理的重要基础模块。页面采用四Tab布局，分别覆盖项目/机组管理、试验对象路径树管理、测量装置台账管理和试验报告导出功能。'
    elif i == 148 and t == '项目基础信息维护；':
        text_by_index[i] = '项目/机组基础信息维护（支持批量CSV导入）；'
    elif i == 150 and t == '试验对象分类管理；':
        text_by_index[i] = '试验对象路径树管理（四级层级：系统→贯穿件→阀门/其他部件）；'
    elif i == 151 and t == '测量装置基础信息管理；':
        text_by_index[i] = '测量装置台账管理（支持按通信方式/启用状态筛选）；'
    elif i == 152 and t == '台账数据统一维护；':
        text_by_index[i] = '试验报告导出（支持Excel/PDF格式，4种导出范围）；'
    elif i == 153 and t == '试验对象与设备关联管理。':
        text_by_index[i] = '试验对象与试验路径配方关联管理。'
    elif i == 154 and '标准化' in t and '结构化' in t:
        text_by_index[i] = '系统通过基础台账管理，实现试验数据标准化、结构化与统一化管理，为后续数据采集、数据分析及历史追溯提供基础支撑。四Tab布局使各类基础数据可在同一页面内快速切换维护。'
    elif i == 157 and '台账管理还支持' in t:
        text_by_index[i] = '台账管理页面还包含独立的测量装置Tab，支持查看装置的基础信息、通信方式、启用状态、连接状态、最近导入时间等。也支持对设备台账进行添加、删除、修改等操作，支持按通信方式和启用状态筛选。报告导出Tab支持Excel和PDF格式，可导出全部记录、本月记录、本月合格或不合格记录。'
    elif i == 160 and p.style.name == 'Heading 3' and t == '实时监控界面':
        # 不改标题，在原位置前面插入新章节
        pass
    elif i == 161 and '实时监控界面主要用于' in t and '实时测量数据' in t:
        text_by_index[i] = '数据分析界面主要用于对数据分析-管理软件中的历史试验数据进行多维度统计分析与可视化展示。系统提供五个分析维度：故障趋势分析、合格率统计、泄漏率趋势、阀门试验次数统计、机组合格情况对比。整个界面采用工业数据分析平台设计风格，结合图表分析与数据统计方式，实现复杂试验数据的可视化展示与快速分析。所有分析结果支持导出为多Sheet的Excel文件。'
    elif i == 162 and '工艺系统、贯穿件、阀门类型' in t:
        text_by_index[i] = '页面顶部设置条件筛选区域，用于实现试验数据的快速查询与精准定位。系统支持按照项目、机组、系统以及试验时间范围等条件进行组合筛选（级联过滤）。用户可通过多条件查询方式快速定位目标试验数据。系统同时支持条件重置功能，满足大型核电项目多系统、多设备环境下的数据检索需求。'
    elif i == 163 and '统计卡片形式' in t:
        text_by_index[i] = '数据分析界面包含五个分析Tab页：'
    elif i == 164 and '图 4 实时监控界面' in t:
        text_by_index[i] = '图 4 数据分析界面'
    elif i == 167 and '试验记录界面主要用于' in t and '测量过程中' in t:
        text_by_index[i] = '试验记录界面主要用于对数据分析-管理软件中的历史试验数据进行集中管理、查询、分析与追溯，是整个数据管理软件中的核心业务模块之一。系统通过试验记录模块，实现试验全过程数字化存档，包括试验数据查看、过程曲线回放、试验路径关联修改、批量操作以及报告导出，为后续的数据分析、设备状态评估、异常追溯以及技术归档提供完整的数据支撑。'
    elif i == 168 and '试验记录界面顶部' in t:
        text_by_index[i] = '试验记录界面顶部设置数据筛选与查询区域，用于对历史试验数据进行快速定位与条件查询。系统支持按照项目、机组、试验结果、时间范围以及关键字（记录编号、试验对象、测量装置、数据包名称）等条件进行组合筛选。记录列表采用分页显示，包含记录编号、项目、机组、对象编码、节点名称、最终泄漏率、泄漏限值、判定结果、关联试验路径、测量装置、操作人员、试验时间、导入时间、备注等字段。'
    elif i == 169 and '试验记录界面具备' in t and '导出与报告' in t:
        text_by_index[i] = '试验记录界面具备过程曲线回放功能。选中一条记录后，底部自动加载该记录的三张趋势曲线图（压力、温度、流量），支持时间范围裁剪、图表交互（拖拽平移、滚轮缩放、悬停查看数值）。右侧面板显示各通道图例（名称、最小-最大值范围）及关联试验路径的完整参数信息。'
    elif i == 171 and '图 5 实验记录' in t:
        text_by_index[i] = '图 5 试验记录界面（含过程曲线回放）'
    elif i == 173 and '系统管理界面主要用于' in t and '用户账号' in t:
        text_by_index[i] = '系统管理界面主要用于实现智能安全壳隔离阀泄漏率测量数据管理软件中的系统管理功能，包含用户权限管理、角色管理、操作日志管理、数据备份与恢复、数据库高可用配置等五个功能Tab页，是整个系统安全管理与运维保障的重要组成部分。'
    elif i == 174 and '系统管理界面顶部' in t:
        text_by_index[i] = '用户权限管理Tab支持用户的增删改查、角色分配、启用/禁用操作。角色管理Tab展示系统内置的三种角色（管理员/试验工程师/只读用户）及其权限范围。操作日志Tab支持按操作类型、时间范围筛选查询，并提供日志预览清理、日志导出和保留天数配置功能。数据备份Tab支持手动备份、数据库还原、备份路径配置、自动备份开关与间隔设置、备份保留策略以及备份历史查看。数据库高可用Tab显示当前主从库连接状态，支持配置主从切换参数和SQL Server连接。'

# 应用文本替换
changed = 0
for i, new_text in text_by_index.items():
    p = doc.paragraphs[i]
    runs = p.runs
    if not runs:
        p.add_run(new_text)
    else:
        for r in runs:
            r.text = ''
        runs[0].text = new_text
    changed += 1
print(f'1. 按索引替换文本：{changed} 处')

# ============== 2. 权限表 ==============
for table in doc.tables:
    rows = table.rows
    if len(rows) >= 4 and norm(rows[0].cells[0].text) == '角色':
        data = [
            ('只读用户(viewer)', '只读查看与导出', '浏览试验记录、数据分析图表、导出试验报告（不可编辑、不可删除）'),
            ('试验工程师(operator)', '试验执行与数据管理', '编辑台账/配方/变量、上传数据、导出数据、修改试验路径关联（不可删除记录、不可进入系统设置）'),
            ('管理员(admin)', '全权管理', '用户管理、角色管理、装置管理、数据备份/恢复、系统配置、删除数据/记录、操作日志管理（需二次确认）'),
        ]
        for ri, (c0, c1, c2) in enumerate(data, start=1):
            for ci, val in enumerate([c0, c1, c2]):
                for para in rows[ri].cells[ci].paragraphs:
                    for r in para.runs: r.text = ''
                    if para.runs: para.runs[0].text = val
                    break
        print('2. 权限表：已更新')
        break

# ============== 3. 图片占位 ==============
img_count = 0
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall('.//' + qn('w:drawing'))
    if not drawings:
        continue
    next_t = norm(doc.paragraphs[i+1].text) if i+1 < len(doc.paragraphs) else ''
    if '图 1' in next_t: ph = '【图片占位：首页概览界面截图】'
    elif '图 2' in next_t: ph = '【图片占位：基础台账管理界面截图（4个Tab）】'
    elif '图 3' in next_t: ph = '【图片占位：测量装置台账管理界面截图】'
    elif '图 4' in next_t: ph = '【图片占位：数据分析界面截图（5个Tab）】'
    elif '图 5' in next_t: ph = '【图片占位：试验记录界面截图（含曲线回放）】'
    elif i < 15: ph = '【图片占位：封面Logo/装饰图】'
    else: ph = '【图片占位：软件主界面概念图（整体布局概览）】'
    for r in p.runs: r.text = ''
    for r_elem in p._element.findall(qn('w:r')):
        if r_elem.findall('.//' + qn('w:drawing')):
            p._element.remove(r_elem)
    if p.runs: p.runs[0].text = ph
    else: p.add_run(ph)
    img_count += 1
print(f'3. 图片占位：{img_count} 处')

# ============== 4. 在"实时监控界面"标题前插入新章节 ==============
# 找到"实时监控界面"标题
rm_title_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 3' and norm(p.text) == '实时监控界面':
        rm_title_idx = i
        break

# 获取格式模板
heading_tmpl = doc.paragraphs[rm_title_idx]  # Heading 3
normal_tmpl = next(p for p in doc.paragraphs if p.style.name == 'Normal' and p.runs)
caption_tmpl = next(p for p in doc.paragraphs if p.style.name == 'Caption')

def make_para(text, tmpl):
    new_p = deepcopy(tmpl._element)
    runs = new_p.findall(qn('w:r'))
    first = True
    for r in runs:
        for t in r.findall(qn('w:t')):
            if first:
                t.text = text
                first = False
            else:
                t.text = ''
    return new_p

rm_elem = doc.paragraphs[rm_title_idx]._element
body = rm_elem.getparent()

# 按正序准备新段落
new_paragraphs = [
    # 试验路径管理页面
    ('试验路径管理页面', heading_tmpl),
    ('试验路径管理页面用于管理试验配方，定义试验参数（泄漏率限值、预充压压力、阀门规格等）。页面顶部提供搜索框、启用状态过滤、CSV导入/导出功能，以及新增、编辑、删除按钮。配方列表以数据表格形式展示，包含试验路径名称、序号、所属系统、贯穿件直径、试验阀门编号、泄漏率限值、预充压P2、创建时间、启用状态等字段。双击表格行或点击编辑按钮可打开配方编辑对话框。', normal_tmpl),
    ('配方编辑对话框分为三个区域：基础信息（名称、序号、系统、启用状态、备注）、阀门参数（贯穿件直径、试验阀门编号、阀门公称直径）、试验参数（泄漏率设计最大值、预充压压力P2）。每次保存编辑时系统自动创建新版本快照（RecipeVersion），支持查看完整的配方修改历史。试验记录中保存的是导入时的配方快照（JSON），后续修改配方不会影响已生成的历史记录。', normal_tmpl),
    ('【图片占位：试验路径管理界面截图】', normal_tmpl),
    ('图 6 试验路径管理界面', caption_tmpl),
    # 实时监视界面
    ('实时监视界面', heading_tmpl),
    ('实时监视界面是软件的核心功能模块，主要用于连接PLC实时采集数据并以趋势曲线形式展示。页面顶部设置试验对象选择区域（项目→机组→试验对象→测量装置，四级级联过滤）和PLC连接控制区域（IP地址输入、连接/断开、开始/停止监视、导出CSV）。系统自动识别Modbus TCP和Siemens S7两种协议，连接成功后启动定时采集（默认1000ms间隔）。', normal_tmpl),
    ('页面中部为实时变量表格，支持在线编辑变量名称、西门子地址、寄存器地址、数据类型、单位、显示范围等属性，并可通过显示开关控制各通道曲线的显隐。页面下部为三张趋势曲线图（压力MPa、温度℃、流量Nml/min），变量按曲线通道属性自动分组到对应图表。图表支持左键拖拽平移、滚轮缩放Y轴、鼠标悬停显示Tracker浮层。通过显示时长设置和自动跟随开关，可灵活控制视口范围。Y轴始终按当前可见窗口内的数据自适应范围。监视过程中连续3次读取失败会触发自动重连。停止监视时系统自动计算最终泄漏率并判定合格/不合格。', normal_tmpl),
    ('【图片占位：实时监视界面截图（含趋势曲线和变量表格）】', normal_tmpl),
    ('图 7 实时监视界面', caption_tmpl),
]

# 逐个插入到 rm_elem 之前
for text, tmpl in new_paragraphs:
    new_elem = make_para(text, tmpl)
    idx = list(body).index(rm_elem)
    body.insert(idx, new_elem)
print(f'4. 新章节：已插入 {len(new_paragraphs)} 个段落')

# ============== 5. 将"实时监控界面"改为"数据分析界面"，图4→图8 ==============
for p in doc.paragraphs:
    t = norm(p.text)
    if p.style.name == 'Heading 3' and t == '实时监控界面':
        # 这是原"实时监控"标题（新插入的实时监视界面已经在它前面了）
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = '数据分析界面'
    elif p.style.name == 'Caption' and t == '图 4 实时监控界面':
        for r in p.runs: r.text = ''
        if p.runs: p.runs[0].text = '图 8 数据分析界面'
print('5. 标题更新：实时监控→数据分析，图4→图8')

# ============== 保存 ==============
doc.save(OUTPUT)
print(f'\n已保存到：{OUTPUT} ({os.path.getsize(OUTPUT)/1024:.1f} KB)')
