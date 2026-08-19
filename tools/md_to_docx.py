# -*- coding: utf-8 -*-
"""
将用户操作手册.md转成Word文档，保持简洁清晰。
"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============ 页面设置 ============
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ============ 样式设置 ============
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '黑体'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        h.font.size = Pt(18)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(15)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(6)

def add_text_with_bold(para, text):
    """处理 **粗体** 标记"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

def add_code_block(text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============ 读取并解析 Markdown ============
with open('doc/用户操作手册.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code_block = False
code_lines = []
in_table = False
table_headers = []
table_rows = []

while i < len(lines):
    line = lines[i].rstrip('\n')

    # 代码块
    if line.strip().startswith('```'):
        if in_code_block:
            add_code_block('\n'.join(code_lines))
            code_lines = []
            in_code_block = False
        else:
            # 先处理可能正在进行的表格
            if in_table:
                add_table(table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # 空行
    if not line.strip():
        if in_table:
            add_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
        i += 1
        continue

    # 分割线
    if line.strip() == '---':
        if in_table:
            add_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
        i += 1
        continue

    # 标题
    if line.startswith('#'):
        if in_table:
            add_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
        level = len(line) - len(line.lstrip('#'))
        text = line.lstrip('#').strip()
        if level <= 3:
            doc.add_heading(text, level=min(level, 3))
        i += 1
        continue

    # 表格
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # 检查是否是分隔行 (|---|---|)
        if all(set(c) <= set('-: ') for c in cells):
            i += 1
            continue
        if not in_table:
            in_table = True
            table_headers = cells
        else:
            table_rows.append(cells)
        i += 1
        continue

    # 列表项
    if re.match(r'^\s*[-*]\s', line) or re.match(r'^\s*\d+\.\s', line):
        if in_table:
            add_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
        text = re.sub(r'^\s*[-*]\s', '', line)
        text = re.sub(r'^\s*\d+\.\s', '', line)
        # 去除 markdown 格式
        text = text.replace('**', '').replace('*', '').replace('`', '')
        p = doc.add_paragraph(text, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)
        i += 1
        continue

    # 普通段落
    if in_table:
        add_table(table_headers, table_rows)
        in_table = False
        table_headers = []
        table_rows = []
    text = line.strip()
    if text.startswith('>'):
        text = text.lstrip('>').strip()
    text = text.replace('`', '').replace('⚠️', '')
    if text:
        p = doc.add_paragraph()
        add_text_with_bold(p, text)
    i += 1

# 收尾
if in_table:
    add_table(table_headers, table_rows)

# ============ 保存 ============
OUTPUT = 'doc/用户操作手册.docx'
doc.save(OUTPUT)
print(f'已保存到：{OUTPUT}')
print(f'文件大小：{os.path.getsize(OUTPUT)/1024:.1f} KB')
