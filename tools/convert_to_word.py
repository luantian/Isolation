from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(run, font_name='宋体', font_size=None, bold=False):
    """设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
    if font_size:
        run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_font(run, font_size=16 if level == 1 else 14, bold=True)
    return heading

def add_paragraph(doc, text, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, font_size=12, bold=bold)
    return p

def add_list_item(doc, text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    set_font(run, font_size=12)
    return p

def add_code_block(doc, code):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code)
    set_font(run, font_name='Consolas', font_size=10)
    return p

def main():
    doc = Document()
    
    # 标题
    add_heading(doc, '智能安全壳隔离阀泄漏率数据管理软件', 1)
    add_heading(doc, '安装部署指南', 1)
    
    # 版本信息
    p = doc.add_paragraph()
    run = p.add_run('版本：v2026.08')
    set_font(run, font_size=12)
    p = doc.add_paragraph()
    run = p.add_run('发布日期：2026 年 8 月')
    set_font(run, font_size=12)
    
    doc.add_paragraph()  # 空行
    
    # 1. 系统要求
    add_heading(doc, '1. 系统要求', 1)
    add_list_item(doc, '操作系统：Windows 10/11（64 位）')
    add_list_item(doc, '数据库：SQL Server 2019 Express 或更高版本')
    add_list_item(doc, '.NET 运行时：已包含在发布包中，无需单独安装')
    
    # 2. 安装 SQL Server Express
    add_heading(doc, '2. 安装 SQL Server Express', 1)
    add_paragraph(doc, '从微软官网下载 SQL Server Express：')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('https://www.microsoft.com/zh-cn/sql-server/sql-server-downloads')
    set_font(run, font_name='Consolas', font_size=10)
    
    add_heading(doc, '2.1 安装 SQL Server', 2)
    add_list_item(doc, '双击下载的 SQLServer*-Express.exe，选择"基本"安装类型')
    add_list_item(doc, '接受许可条款，点击"下一步"')
    add_list_item(doc, '选择实例：保持默认的命名实例 SQLEXPRESS，不要修改')
    add_list_item(doc, '安装路径：保持默认即可，点击"下一步"开始安装')
    add_list_item(doc, '等待安装完成，点击"关闭"')
    
    add_paragraph(doc, '建议同时安装 SQL Server Management Studio (SSMS)：')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('https://docs.microsoft.com/zh-cn/sql/ssms/download-sql-server-management-studio-ssms')
    set_font(run, font_name='Consolas', font_size=10)
    
    add_heading(doc, '2.2 配置数据库', 2)
    add_paragraph(doc, '安装完成后需要启用 sa 账户并配置网络访问，按以下步骤操作：')
    
    add_paragraph(doc, '第一步：设置 sa 账户', bold=True)
    add_list_item(doc, '打开 SSMS，在"连接到服务器"对话框中：')
    add_list_item(doc, '服务器名称：localhost\\SQLEXPRESS（或 .\\SQLEXPRESS）', level=1)
    add_list_item(doc, '身份验证：Windows 身份验证', level=1)
    add_list_item(doc, '点击"连接"', level=1)
    add_list_item(doc, '在左侧"对象资源管理器"中，展开"安全性" → "登录名"')
    add_list_item(doc, '双击 sa，在弹出窗口中：')
    add_list_item(doc, '选择"常规"页 → 设置新密码（请记住此密码）', level=1)
    add_list_item(doc, '选择"状态"页 → 将"登录"设置为启用', level=1)
    add_list_item(doc, '点击"确定"', level=1)
    
    add_paragraph(doc, '第二步：启用混合模式和网络访问', bold=True)
    add_list_item(doc, '右键点击服务器名称（最顶部节点）→ "属性" → "安全性"')
    add_list_item(doc, '将"服务器身份验证"改为 SQL Server 和 Windows 身份验证模式', level=1)
    add_list_item(doc, '点击"确定"', level=1)
    add_list_item(doc, '打开 SQL Server Configuration Manager（开始菜单搜索即可）')
    add_list_item(doc, '展开"SQL Server 网络配置" → 点击"SQLEXPRESS 的协议"')
    add_list_item(doc, '右键"TCP/IP" → "启用"', level=1)
    
    add_paragraph(doc, '第三步：重启服务', bold=True)
    add_list_item(doc, '按 Win + R，输入 services.msc')
    add_list_item(doc, '找到 SQL Server (SQLEXPRESS)，右键 → "重新启动"')
    
    add_heading(doc, '2.3 验证 sa 账户登录', 2)
    add_list_item(doc, '在 SSMS 中"断开连接"')
    add_list_item(doc, '重新"连接"，在对话框中：')
    add_list_item(doc, '服务器名称：localhost\\SQLEXPRESS', level=1)
    add_list_item(doc, '身份验证：SQL Server 身份验证', level=1)
    add_list_item(doc, '登录名：sa', level=1)
    add_list_item(doc, '密码：刚才设置的密码', level=1)
    add_list_item(doc, '点击"连接"，验证是否成功')
    
    # 3. 安装管理软件
    add_heading(doc, '3. 安装管理软件', 1)
    add_paragraph(doc, '将 IsolationLeakageApp_v20260803_win-x64_SelfContained.zip 解压到目标目录（如 C:\\IsolationLeakage\\）。')
    add_paragraph(doc, '主程序：IsolationLeakage.App.exe')
    
    # 4. 配置数据库连接
    add_heading(doc, '4. 配置数据库连接', 1)
    add_paragraph(doc, '编辑软件目录下的 appsettings.json，修改 DefaultConnection：')
    add_code_block(doc, '{\n  "ConnectionStrings": {\n    "DefaultConnection": "Server=localhost\\\\SQLEXPRESS;Database=IsolationLeakageDb;User Id=sa;Password=YourPassword;Connect Timeout=10;Trust Server Certificate=True;",\n    "SecondaryConnection": ""\n  }\n}')
    
    add_paragraph(doc, '参数说明：', bold=True)
    add_list_item(doc, 'Server：数据库服务器地址（单机用 localhost，远程用 IP）')
    add_list_item(doc, 'Database：固定为 IsolationLeakageDb')
    add_list_item(doc, 'User Id / Password：数据库账户密码')
    
    p = doc.add_paragraph()
    run = p.add_run('注意：无需手动创建数据库。软件首次启动时会自动创建 IsolationLeakageDb 并初始化表结构。')
    set_font(run, font_size=12, bold=True)
    
    # 5. 启动与验证
    add_heading(doc, '5. 启动与验证', 1)
    add_paragraph(doc, '双击 IsolationLeakage.App.exe 启动软件。')
    
    add_paragraph(doc, '默认账户：', bold=True)
    add_list_item(doc, '管理员：admin / admin123')
    add_list_item(doc, '操作员：operator / operator123')
    add_list_item(doc, '只读用户：viewer / viewer123')
    add_paragraph(doc, '首次登录后请立即修改密码。')
    
    add_paragraph(doc, '验证步骤：', bold=True)
    add_list_item(doc, '查看左下角数据库状态（● 绿色表示正常）')
    add_list_item(doc, '尝试登录系统')
    add_list_item(doc, '进入"系统设置"确认连接正常')
    
    # 6. 常见问题
    add_heading(doc, '6. 常见问题', 1)
    
    add_paragraph(doc, 'Q: 提示"数据库连接失败"？', bold=True)
    add_list_item(doc, '检查 SQL Server 服务是否启动（services.msc → SQL Server (SQLEXPRESS)）')
    add_list_item(doc, '核对 appsettings.json 中的服务器地址、用户名、密码')
    add_list_item(doc, '远程部署时检查防火墙是否开放 1433 端口')
    
    add_paragraph(doc, 'Q: 提示"数据库迁移失败"？', bold=True)
    add_list_item(doc, '确保 sa 用户具有 db_owner 角色')
    add_list_item(doc, '检查数据库文件目录是否有写入权限')
    
    add_paragraph(doc, 'Q: 日志文件在哪里？', bold=True)
    add_list_item(doc, '位于软件目录下的 logs\\ 文件夹')
    add_list_item(doc, '遇到问题可将日志发给技术支持')
    
    # 附录
    add_heading(doc, '附录：技术支持', 1)
    add_paragraph(doc, '遇到问题请提供：')
    add_list_item(doc, '软件版本号（登录界面右下角）')
    add_list_item(doc, '操作系统版本')
    add_list_item(doc, '错误截图或 logs\\ 目录下的日志文件')
    
    # 保存文档
    output_path = r'F:\workspace\cechuang\projects\Isolation\doc\客户安装部署指南_v2.docx'
    doc.save(output_path)
    print(f'Word 文档已生成：{output_path}')

if __name__ == '__main__':
    main()
