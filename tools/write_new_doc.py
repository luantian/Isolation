# -*- coding: utf-8 -*-
"""
从 A(1) 原始文档出发，按照正确结构写全新文件 A(4)。
结构完全按代码中的导航顺序，核心3页面展开，其他精简。
"""
import sys, os, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

INPUT = 'doc/数据分析-管理软件设计方案-A(1).docx'
OUTPUT = 'doc/数据分析-管理软件设计方案-A(4).docx'

doc = Document(INPUT)

# ============ 工具函数 ============
def norm(t): return t.replace('\xa0', ' ').strip()

def set_text(para, new_text):
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    for r in runs: r.text = ''
    runs[0].text = new_text

def make_para(text, template_para):
    new_p = deepcopy(template_para._element)
    runs = new_p.findall(qn('w:r'))
    first = True
    for r in runs:
        for t in r.findall(qn('w:t')):
            if first:
                t.text = text
                first = False
            else:
                t.text = ''
    return new_p

def insert_before(target_elem, new_elem):
    parent = target_elem.getparent()
    idx = list(parent).index(target_elem)
    parent.insert(idx, new_elem)

def remove_para(p):
    p._element.getparent().remove(p._element)

# ============ 第一步：修改现有段落文本 ============
print('=== 第一步：修改现有段落文本 ===')
text_fixes = {}
for i, p in enumerate(doc.paragraphs):
    t = norm(p.text)
    # 工作范围
    if i == 43 and '自动上传' in t:
        text_fixes[i] = '实现试验数据导入（单文件与批量）、自动归档、自动分类存储；'
    # 通信层
    elif i == 51 and 'RS485' in t and 'Siemens' not in t:
        text_fixes[i] = '通信层负责实现RS232、RS485、TCP/IP、Modbus TCP、Siemens S7等协议通信；'
    # 应用层
    elif i == 53 and 'Web管理界面' in t:
        text_fixes[i] = '应用层负责向用户提供桌面管理界面、数据分析界面、报表界面以及系统配置界面。'
    # 模块名 [58-63]
    elif i == 58 and '主控台' in t:
        text_fixes[i] = '首页概览：系统的全局数据总览仪表盘，展示核心KPI指标、试验记录预览和系统运行状态。'
    elif i == 59 and '资产管理中心' in t:
        text_fixes[i] = '试验对象：管理项目/机组、试验对象路径树（四级层级）、测量装置台账，以及试验报告导出。'
    elif i == 60 and '任务管理中心' in t:
        text_fixes[i] = '试验路径：管理试验配方，定义泄漏率限值、预充压压力、阀门规格等试验参数。'
    elif i == 61 and '数据中心' in t and len(t) < 80:
        text_fixes[i] = '试验记录：存储、查询和管理所有历史试验记录，支持过程曲线回放与批量操作。'
    elif i == 62 and '分析中心' in t:
        text_fixes[i] = '实时监视：连接PLC实时采集压力、温度、流量数据，以趋势曲线形式展示。'
    elif i == 63 and '系统设置' in t and '用户管理' in t:
        text_fixes[i] = '数据分析：提供故障趋势、合格率统计、泄漏率趋势等多维度统计分析。'
    # 服务器配置
    elif i == 65 and '8核CPU' in t:
        text_fixes[i] = '服务器配置：配置不低于4核CPU / 8GB内存 / 256GB SSD存储，操作系统采用 Windows 10/11 或 Windows Server 2016+ 或同等稳定的Windows桌面/服务器版本。'
    # 数据上传
    elif i == 78 and '双向数据交互' in t:
        text_fixes[i] = '该模块主要实现数据管理软件与多台测量装置之间的双向数据交互，支持单文件导入与批量导入，主要功能包含：'
    elif i == 79 and '试验对象下载' in t:
        text_fixes[i] = '试验任务下发：软件可将选定的试验对象（含配置参数）下发至指定的测量装置。'
    elif i == 80 and '试验数据上传' in t:
        text_fixes[i] = '试验数据上传：支持单文件导入（.json/.txt/.csv）和批量导入（选择文件夹自动解析匹配），系统自动识别试验对象并关联试验路径配方。'
    elif i == 81 and '非覆盖存储' in t:
        text_fixes[i] = '非覆盖存储：同一试验对象的不同次试验按时间顺序追加存储，每次导入自动保存配方快照（JSON），确保历史数据完整。'
    # 数据分析
    elif i == 83 and '主要功能如下' in t:
        text_fixes[i] = '本模块将原始试验数据转化为统计图表与趋势分析结果，提供五个分析维度：故障趋势（按阀门类型统计合格/不合格数）、合格率统计（各阀门合格率评定）、泄漏率趋势（多系列曲线对比）、阀门试验次数（Top 50排名）、机组合格情况（各机组合格率对比）。支持按项目/机组/系统/时间范围筛选，可导出多Sheet的Excel文件。'
    elif i == 84 and t == '单阀门分析：':
        text_fixes[i] = '故障趋势分析：'
    elif i == 85 and '历史泄漏率' in t:
        text_fixes[i] = '按阀门类型统计合格数与不合格数，以堆叠柱状图展示故障分布。'
    elif i == 86 and '试验合格/失败' in t:
        text_fixes[i] = '展示总试验数、总体合格率、不合格数，并计算各阀门的合格率（≥95%合格，80%-95%注意，<80%不合格）。'
    elif i == 87 and t == '机组级分析：':
        text_fixes[i] = '泄漏率趋势分析：'
    elif i == 88 and '按阀门类型' in t:
        text_fixes[i] = '以多系列趋势曲线展示不同阀门类型的泄漏率历史变化。'
    elif i == 89 and '合格率仪表盘' in t:
        text_fixes[i] = '阀门试验次数统计：排名展示各阀门的试验次数（Top 50）。'
    elif i == 90 and '多因素' in t:
        text_fixes[i] = '机组合格情况统计：按机组展示合格率，以柱状图对比各机组。所有分析结果支持导出为Excel文件。'
    # 权限
    elif i == 94 and '验证码' in t:
        text_fixes[i] = '认证：用户名+密码登录，密码加密存储，支持会话管理。'
    elif i == 95 and '审计' in t:
        text_fixes[i] = '审计：所有关键操作记录日志，支持按时间和级别筛选，支持日志清理与导出。'
    elif i == 96 and '数据隔离' in t:
        text_fixes[i] = '数据隔离：通过角色权限控制不同用户的可见页面和可执行操作。'
    # 备份
    elif i == 98 and '双服务器' in t:
        text_fixes[i] = '本模块通过自动/手动备份、主从数据库自动故障切换及灵活的数据恢复机制，确保数据安全与业务连续性。'
    elif i == 99 and '每日/每周' in t:
        text_fixes[i] = '自动备份：支持可配置间隔定时全量备份，可配置保留天数。'
    elif i == 101 and '实时同步' in t:
        text_fixes[i] = '主从切换：主库故障时自动切换至从库，恢复后自动切回，切换期间通过磁盘缓冲确保数据零丢失。'
    # 技术组件
    elif i == 113 and '二进制日志' in t:
        text_fixes[i] = '文件操作：实现本地试验数据文件的导入（JSON/CSV/TXT数据包）、导出（报告/原始数据），以及配置文件的读取与保存。'
    elif i == 114 and 'HTTP/HTTPS' in t:
        text_fixes[i] = '数据对接：封装与测量装置的通信协议（Modbus TCP、Siemens S7、RS232/RS485、USB），负责实时采集PLC数据、下发试验任务、接收上传数据，提供数据缓存与断点续传支持。'

count = 0
for i, new_text in text_fixes.items():
    set_text(doc.paragraphs[i], new_text)
    count += 1
print(f'修改了 {count} 段文本')

# ============ 第二步：删除"数据分析与可视化"子节（内容已合并到模块描述和界面章节）============
print('\n=== 第二步：删除冗余章节 ===')
# 找到"数据分析与可视化"Heading 2
analysis_start = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and '数据分析与可视化' in p.text:
        analysis_start = i
        break

if analysis_start:
    # 找到下一个 Heading 2
    analysis_end = None
    for j in range(analysis_start + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 2':
            analysis_end = j
            break
    if analysis_end is None:
        analysis_end = len(doc.paragraphs)
    # 删除
    for j in range(analysis_end - 1, analysis_start - 1, -1):
        remove_para(doc.paragraphs[j])
    print(f'删除"数据分析与可视化"：{analysis_end - analysis_start} 段')

# ============ 第三步：重写界面章节 ============
print('\n=== 第三步：重写界面章节 ===')

# 获取模板
h3_tmpl = list_tmpl = normal_tmpl = caption_tmpl = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 3' and not h3_tmpl: h3_tmpl = p
    if 'List' in (p.style.name or '') and p.runs and not list_tmpl: list_tmpl = p
    if p.style.name == 'Normal' and p.runs and not normal_tmpl: normal_tmpl = p
    if p.style.name == 'Caption' and not caption_tmpl: caption_tmpl = p

# 找到"软件主界面概念图详解"之后的所有 Heading 3 段落
section_start = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and '概念图详解' in p.text:
        section_start = i + 1
        break

# 找到"系统联调方案" Heading 1
section_end = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and '系统联调' in p.text:
        section_end = i
        break

print(f'界面章节范围：[{section_start}] - [{section_end}]')

# 删除旧界面章节内容
for j in range(section_end - 1, section_start - 1, -1):
    remove_para(doc.paragraphs[j])
print('已删除旧界面章节内容')

# 重新定位插入点
insert_before_elem = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and '系统联调' in p.text:
        insert_before_elem = p._element
        break

body = insert_before_elem.getparent()

# ============ 新的界面章节内容 ============
sections = []

# --- 1. 首页概览（精简）---
sections.append(('h3', '首页概览'))
sections.append(('normal', '首页概览采用"左侧导航 + 顶部核心指标 + 中部业务监控 + 底部台账概况"布局，以6个KPI卡片（试验对象数、测量装置数、历史记录数、本月合格率、待处理异常数、最近备份时间）、试验记录预览表、最近导入详情、9项台账统计指标和系统维护状态为主体，满足全局数据一屏总览的需求。'))
sections.append(('img', '【图片占位：首页概览界面截图】'))
sections.append(('cap', '图 1 首页概览界面'))

# --- 2. 试验对象页面（⭐重点展开）---
sections.append(('h3', '试验对象页面'))
sections.append(('normal', '试验对象页面是系统的基础数据管理核心，采用四Tab布局，分别覆盖项目/机组管理、试验对象路径树管理、测量装置台账管理和试验报告导出功能。'))

sections.append(('normal', 'Tab 1 - 项目/机组：系统的最顶层组织单元。项目用于归类和组织所有试验数据，机组隶属于项目。支持项目和机组的增删改查操作（项目编码自动生成，格式P{年月}{序号}），支持通过CSV文件批量导入。左侧为项目列表，右侧为机组列表（按选中项目过滤）。'))

sections.append(('normal', 'Tab 2 - 试验对象管理：采用四级层级树结构——系统(System) → 贯穿件(Penetration) → 阀门(Valve)/其他部件(OtherComponent)。页面顶部提供项目/机组下拉框范围过滤和关键字搜索。左侧为可展开折叠的试验对象树，底部提供四个快速新建按钮。右侧显示选中节点的详细信息（编号、名称、类型、泄漏率限值、试验压力、备注）及操作按钮（修改、导入、导出、删除），下方展示该对象的试验统计和关联的试验路径配方信息。叶子节点可配置默认关联试验路径，用于数据导入时自动匹配配方。'))

sections.append(('normal', 'Tab 3 - 测量装置：以表格形式展示所有测量装置信息（编号、名称、IP、序列号、通信方式、启用状态、连接状态、同步时间、导入时间）。支持按通信方式和启用状态筛选。提供装置的新增、编辑、删除操作，编号自动生成（格式DEV-{时间戳}）。只有启用状态的装置才能在实时监视中被选择。'))

sections.append(('normal', 'Tab 4 - 报告导出：支持导出全部/本月/本月合格/本月不合格四种范围的试验记录，支持Excel和PDF两种格式。可自定义文件名和导出目录。提供快速导出按钮一键生成报告。'))

sections.append(('img', '【图片占位：试验对象管理界面截图（4个Tab）】'))
sections.append(('cap', '图 2 试验对象管理界面'))

# --- 3. 试验路径页面（精简）---
sections.append(('h3', '试验路径页面'))
sections.append(('normal', '试验路径页面用于管理试验配方，定义泄漏率限值、预充压压力、阀门规格等试验参数。页面提供搜索、启用过滤、CSV导入导出、以及配方的增删改功能。配方列表展示名称、序号、系统、贯穿件直径、阀门编号、泄漏率限值、预充压P2、启用状态等字段。'))
sections.append(('normal', '配方编辑对话框分三个区域：基础信息（名称、系统、启用状态、备注）、阀门参数（贯穿件直径、阀门编号、公称直径）、试验参数（泄漏率设计最大值、预充压P2）。每次编辑自动创建版本快照（RecipeVersion），试验记录中保存导入时的配方快照（JSON），后续修改配方不影响历史记录。'))
sections.append(('img', '【图片占位：试验路径管理界面截图】'))
sections.append(('cap', '图 3 试验路径管理界面'))

# --- 4. 试验记录页面（⭐重点展开）---
sections.append(('h3', '试验记录页面'))
sections.append(('normal', '试验记录页面是系统的核心业务模块，用于试验全过程数据的集中管理、查询、分析与追溯。'))

sections.append(('normal', '查询区域支持按项目（级联过滤机组）、试验结果（全部/合格/不合格/未知）、时间范围、关键字（记录编号/试验对象/测量装置/数据包名称）等条件组合筛选。查询结果以分页表格展示，包含记录编号、项目、机组、对象编码、节点名称、最终泄漏率(Nml/min)、泄漏限值(Nml/min)、判定结果（合格绿/不合格红）、关联试验路径、测量装置、操作人员、试验时间、导入时间、备注等字段。支持表头全选批量选择。'))

sections.append(('normal', '过程曲线回放功能位于页面底部。选中记录后自动加载三张趋势曲线图（压力MPa、温度℃、流量Nml/min），曲线按通道属性自动分组。支持左键拖拽平移、滚轮缩放Y轴、鼠标悬停显示Tracker浮层。提供时间范围输入框（0=全部）裁剪显示。右侧面板显示通道图例和关联试验路径的完整参数。'))

sections.append(('normal', '支持批量修改试验路径关联（自动重算合格判定）、批量删除（二次确认）、双击编辑单条记录。数据上传支持单文件导入（.json/.txt/.csv，自动识别对象并关联默认配方）和批量导入（选择文件夹自动解析）。导出支持Excel、PDF、CSV格式。'))

sections.append(('img', '【图片占位：试验记录界面截图（含曲线回放）】'))
sections.append(('cap', '图 4 试验记录界面'))

# --- 5. 实时监视界面（⭐重点展开）---
sections.append(('h3', '实时监视界面'))
sections.append(('normal', '实时监视界面是系统的核心功能模块，用于连接PLC实时采集压力、温度、流量等数据并以趋势曲线展示。'))

sections.append(('normal', '顶部控制区包含四级级联选择（项目→机组→试验对象→测量装置）和PLC连接控制（IP地址、连接/断开、保存地址）。系统自动识别Modbus TCP和Siemens S7两种协议。连接成功后可开始/停止监视、导出CSV。状态以颜色指示（红=未连接，绿=已连接）。连续3次读取失败自动重连，重连失败则停止监视。'))

sections.append(('normal', '中部实时变量表格支持在线编辑：显示开关、颜色、变量名称、西门子地址、寄存器地址、数据类型（double/int/float/real/ushort/dword）、单位、最小/最大值、当前值（只读）、更新时间（只读）、状态（只读）。修改后点"保存配置"持久化到数据库。'))

sections.append(('normal', '底部三张趋势曲线图按通道属性自动分组：压力(MPa)、温度(℃)、流量(Nml/min)。图表支持左键拖拽平移、滚轮缩放、悬停Tracker。通过"显示时长"输入框（默认600秒）和"自动"复选框控制视口——勾选自动时跟随最新数据滚动，Y轴按可见窗口自适应；取消时停在当前位置。所有数据始终保留不裁剪。停止监视时自动计算最终泄漏率（取M1/M2最大值）并判定合格/不合格。'))

sections.append(('normal', '数据安全：周期自动保存（间隔随数据量调整10s→30s→60s→5min），内存缓冲区上限86400点，超出自动裁剪。停止/关闭时同步保存最终版数据。'))

sections.append(('img', '【图片占位：实时监视界面截图（含趋势曲线和变量表格）】'))
sections.append(('cap', '图 5 实时监视界面'))

# --- 6. 数据分析界面（精简）---
sections.append(('h3', '数据分析界面'))
sections.append(('normal', '数据分析界面提供五个分析维度：故障趋势（按阀门类型统计合格/不合格数）、合格率统计（各阀门合格率及等级评定）、泄漏率趋势（多系列曲线对比）、阀门试验次数（Top 50排名）、机组合格情况（各机组合格率对比）。支持按项目/机组/系统/时间筛选，可导出多Sheet的Excel文件。'))
sections.append(('img', '【图片占位：数据分析界面截图（5个Tab）】'))
sections.append(('cap', '图 6 数据分析界面'))

# --- 7. 系统设置页面（精简）---
sections.append(('h3', '系统设置页面'))
sections.append(('normal', '系统设置页面包含五个功能Tab：用户权限（增删改查、角色分配、启用/禁用）、角色管理（管理员/试验工程师/只读用户三种内置角色）、操作日志（按类型/时间筛选、清理与导出、保留天数配置）、数据备份（手动备份、还原、自动备份间隔与保留策略、历史查看）、数据库高可用（主从连接状态、故障切换参数配置、SQL Server连接配置）。'))
sections.append(('img', '【图片占位：系统设置界面截图（5个Tab）】'))
sections.append(('cap', '图 7 系统设置界面'))

# 按倒序插入
for style_type, text in reversed(sections):
    if style_type == 'h3':
        elem = make_para(text, h3_tmpl)
    elif style_type == 'normal':
        elem = make_para(text, normal_tmpl)
    elif style_type == 'list':
        elem = make_para(text, list_tmpl)
    elif style_type == 'img':
        elem = make_para(text, normal_tmpl)
    elif style_type == 'cap':
        elem = make_para(text, caption_tmpl)
    insert_before(insert_before_elem, elem)

print(f'已插入 {len(sections)} 个新段落')

# ============ 第四步：处理图片占位 ============
print('\n=== 第四步：处理图片占位 ===')
img_count = 0
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall('.//' + qn('w:drawing'))
    if not drawings:
        continue
    next_t = norm(doc.paragraphs[i+1].text) if i+1 < len(doc.paragraphs) else ''
    if '图 1' in next_t: ph = '【图片占位：首页概览界面截图】'
    elif '图 2' in next_t: ph = '【图片占位：试验对象管理界面截图（4个Tab）】'
    elif '图 3' in next_t: ph = '【图片占位：试验路径管理界面截图】'
    elif i < 15: ph = '【图片占位：封面Logo/装饰图】'
    else: ph = '【图片占位：软件主界面概念图】'
    for r in p.runs: r.text = ''
    for r_elem in p._element.findall(qn('w:r')):
        if r_elem.findall('.//' + qn('w:drawing')):
            p._element.remove(r_elem)
    if p.runs: p.runs[0].text = ph
    else: p.add_run(ph)
    img_count += 1
print(f'处理了 {img_count} 处图片')

# ============ 第五步：修改权限表 ============
print('\n=== 第五步：修改权限表 ===')
for table in doc.tables:
    rows = table.rows
    if len(rows) >= 4 and norm(rows[0].cells[0].text) == '角色':
        data = [
            ('只读用户(viewer)', '只读查看与导出', '浏览试验记录、数据分析图表、导出试验报告'),
            ('试验工程师(operator)', '试验执行与数据管理', '编辑台账/配方/变量、上传数据、导出数据'),
            ('管理员(admin)', '全权管理', '用户/角色管理、装置管理、备份/恢复、删除数据'),
        ]
        for ri, (c0, c1, c2) in enumerate(data, start=1):
            for ci, val in enumerate([c0, c1, c2]):
                for para in rows[ri].cells[ci].paragraphs:
                    for r in para.runs: r.text = ''
                    if para.runs: para.runs[0].text = val
                    break
        print('权限表已更新')
        break

# ============ 第六步：修改首页描述段落 ============
print('\n=== 第六步：修改首页相关段落 ===')
for i, p in enumerate(doc.paragraphs):
    t = norm(p.text)
    if i == 127 and '主要用于：' in t and '工业监控' in t:
        set_text(p, '软件首页整体采用工业监控平台风格设计，以核心指标卡片、数据表格、台账概况和系统维护四大模块为主体，主要用于：')
    elif i == 128 and t == '泄漏率试验数据管理':
        set_text(p, '泄漏率试验数据管理总览')
    elif i == 129 and t == '多设备在线监控':
        set_text(p, '多设备状态监控')
    elif i == 130 and t == '阀门试验结果分析':
        set_text(p, '试验合格率统计')
    elif i == 133 and '所有设备的运行状态' in t:
        set_text(p, '此区域主要负责全局数据概览，在程序首页即可查看6个核心指标（试验对象数量、测量装置数量、历史记录条数、本月合格率、待处理异常数、最近备份时间）以及最近导入详情。')
    elif i == 134 and '左侧导航 + 中部业务监控' in t:
        set_text(p, '系统首页采用"左侧导航 + 顶部核心指标 + 中部业务监控 + 底部台账概况"的布局结构。')
    elif i == 136 and t == '泄漏率试验全过程数字化管理；':
        set_text(p, '泄漏率试验关键指标一屏总览；')
    elif i == 137 and t == '多设备集中运行监控；':
        set_text(p, '试验记录预览与最近导入详情；')
    elif i == 138 and t == '自动化试验数据采集；':
        set_text(p, '台账概况（项目/机组/系统/贯穿件/阀门/部件数量统计）；')
    elif i == 139 and t == '历史数据统一管理；':
        set_text(p, '系统维护状态（数据库连接、备份状态、同步状态）；')
    elif i == 140 and t == '设备运行状态实时监控；':
        set_text(p, '核心KPI卡片（试验对象数、装置数、记录数、合格率、异常数、备份状态）；')
    elif i == 141 and t == '数据分析与趋势研究。':
        set_text(p, '数据分析与趋势研究快速入口。')
    # 基础台账相关
    elif i == 146 and '基础台账页面用于' in t:
        set_text(p, '试验对象页面用于实现数据分析-管理软件中的基础数据管理功能。页面采用四Tab布局，分别覆盖项目/机组管理、试验对象路径树管理、测量装置台账管理和试验报告导出功能。')
    elif i == 148 and t == '项目基础信息维护；':
        set_text(p, '项目/机组基础信息维护（支持批量CSV导入）；')
    elif i == 150 and t == '试验对象分类管理；':
        set_text(p, '试验对象路径树管理（四级层级：系统→贯穿件→阀门/其他部件）；')
    elif i == 151 and t == '测量装置基础信息管理；':
        set_text(p, '测量装置台账管理（支持按通信方式/启用状态筛选）；')
    elif i == 152 and t == '台账数据统一维护；':
        set_text(p, '试验报告导出（支持Excel/PDF格式，4种导出范围）；')
    elif i == 153 and t == '试验对象与设备关联管理。':
        set_text(p, '试验对象与试验路径配方关联管理。')
    elif i == 154 and '标准化' in t:
        set_text(p, '系统通过试验对象管理，实现试验数据标准化、结构化与统一化管理。四Tab布局使各类基础数据可在同一页面内快速切换维护。')
    elif i == 157 and '台账管理还支持' in t:
        set_text(p, '测量装置Tab支持查看装置基础信息、通信方式、启用状态、连接状态、最近导入时间等，支持添加、删除、修改操作。报告导出Tab支持Excel和PDF格式。')
    # 旧"实时监控界面"描述 → 改为数据分析
    elif i == 161 and '实时监控界面主要用于' in t:
        set_text(p, '数据分析界面提供五个维度的统计分析：故障趋势、合格率统计、泄漏率趋势、阀门试验次数、机组合格情况。支持多条件筛选和Excel导出。')
    elif i == 162 and '工艺系统' in t:
        set_text(p, '页面顶部提供项目/机组/系统/时间范围筛选条件，支持级联过滤和条件重置。')
    elif i == 163 and '统计卡片形式' in t:
        set_text(p, '界面以统计卡片和图表方式展示关键指标。')
    elif i == 164 and '图 4 实时监控' in t:
        set_text(p, '图 6 数据分析界面')
    # 试验记录
    elif i == 167 and '试验记录界面主要用于' in t:
        set_text(p, '试验记录界面是核心业务模块，用于试验全过程数据的集中管理、查询、分析与追溯。')
    elif i == 168 and '试验记录界面顶部' in t:
        set_text(p, '支持按项目/机组/结果/时间/关键字多条件筛选，记录分页显示，含曲线回放和批量操作。')
    elif i == 169 and '试验记录界面具备' in t:
        set_text(p, '支持Excel/PDF/CSV格式导出，可生成标准化试验报告。')
    elif i == 171 and '图 5 实验记录' in t:
        set_text(p, '图 4 试验记录界面')
    # 系统管理
    elif i == 173 and '系统管理界面主要用于' in t:
        set_text(p, '系统管理界面包含五个功能Tab：用户权限、角色管理、操作日志、数据备份、数据库高可用。')
    elif i == 174 and '系统管理界面顶部' in t:
        set_text(p, '支持用户增删改查、角色分配、日志查询清理、备份还原、主从切换配置等操作。')

# ============ 保存 ============
print('\n=== 保存 ===')
tmp = OUTPUT.replace('.docx', '_tmp.docx')
doc.save(tmp)
try: os.remove(OUTPUT)
except: pass
os.rename(tmp, OUTPUT)
print(f'已保存到：{OUTPUT}')
print(f'文件大小：{os.path.getsize(OUTPUT)/1024:.1f} KB')
