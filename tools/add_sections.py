# -*- coding: utf-8 -*-
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

doc = Document('doc/数据分析-管理软件设计方案-A(2).docx')

# 找到"数据分析界面"标题段落
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 3' and p.text.strip() == '数据分析界面':
        target_idx = i
        break
print(f'数据分析界面标题在段落 [{target_idx}]')

# 获取格式模板
heading_template = None
normal_template = None
caption_template = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 3' and not heading_template:
        heading_template = p
    if p.style.name == 'Normal' and p.runs and not normal_template:
        normal_template = p
    if p.style.name == 'Caption' and not caption_template:
        caption_template = p

data_elem = doc.paragraphs[target_idx]._element
parent = data_elem.getparent()

def make_heading(text):
    elem = doc.paragraphs[target_idx]._element.__class__(qn('w:p'))
    src_pPr = heading_template._element.find(qn('w:pPr'))
    if src_pPr is not None:
        elem.append(deepcopy(src_pPr))
    run = elem.makeelement(qn('w:r'), {})
    if heading_template.runs:
        src_rPr = heading_template.runs[0]._element.find(qn('w:rPr'))
        if src_rPr is not None:
            run.append(deepcopy(src_rPr))
    t = run.makeelement(qn('w:t'), {})
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    elem.append(run)
    return elem

def make_body(text):
    elem = doc.paragraphs[target_idx]._element.__class__(qn('w:p'))
    src_pPr = normal_template._element.find(qn('w:pPr'))
    if src_pPr is not None:
        elem.append(deepcopy(src_pPr))
    run = elem.makeelement(qn('w:r'), {})
    src_rPr = normal_template.runs[0]._element.find(qn('w:rPr'))
    if src_rPr is not None:
        run.append(deepcopy(src_rPr))
    t = run.makeelement(qn('w:t'), {})
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    elem.append(run)
    return elem

def make_img_placeholder(text):
    return make_body(text)

def make_caption(text):
    elem = doc.paragraphs[target_idx]._element.__class__(qn('w:p'))
    src_pPr = caption_template._element.find(qn('w:pPr'))
    if src_pPr is not None:
        elem.append(deepcopy(src_pPr))
    run = elem.makeelement(qn('w:r'), {})
    if caption_template.runs:
        src_rPr = caption_template.runs[0]._element.find(qn('w:rPr'))
        if src_rPr is not None:
            run.append(deepcopy(src_rPr))
    t = run.makeelement(qn('w:t'), {})
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    elem.append(run)
    return elem

# === 插入试验路径管理页面 ===
path_elements = [
    make_heading("试验路径管理页面"),
    make_body(
        "试验路径管理页面用于管理试验配方，定义试验参数（泄漏率限值、预充压压力、阀门规格等）。"
        "页面顶部提供搜索框、启用状态过滤、CSV导入/导出功能，以及新增、编辑、删除按钮。"
        "配方列表以数据表格形式展示，包含试验路径名称、序号、所属系统、贯穿件直径、试验阀门编号、"
        "泄漏率限值、预充压P2、创建时间、启用状态等字段。双击表格行或点击编辑按钮可打开配方编辑对话框。"
    ),
    make_body(
        "配方编辑对话框分为三个区域：基础信息（名称、序号、系统、启用状态、备注）、"
        "阀门参数（贯穿件直径、试验阀门编号、阀门公称直径）、试验参数（泄漏率设计最大值、预充压压力P2）。"
        "每次保存编辑时系统自动创建新版本快照（RecipeVersion），支持查看完整的配方修改历史。"
        "试验记录中保存的是导入时的配方快照（JSON），后续修改配方不会影响已生成的历史记录。"
    ),
    make_img_placeholder("【图片占位：试验路径管理界面截图】"),
    make_caption("图 6 试验路径管理界面"),
]

for elem in path_elements:
    parent.insert(list(parent).index(data_elem), elem)
print("已插入：试验路径管理页面章节")

# === 插入实时监视界面 ===
# 重新定位数据分析界面
target_idx2 = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 3' and p.text.strip() == '数据分析界面':
        target_idx2 = i
        break

data_elem2 = doc.paragraphs[target_idx2]._element

rm_elements = [
    make_heading("实时监视界面"),
    make_body(
        "实时监视界面是软件的核心功能模块，主要用于连接PLC实时采集数据并以趋势曲线形式展示。"
        "页面顶部设置试验对象选择区域（项目-机组-试验对象-测量装置，四级级联过滤）和PLC连接控制区域"
        "（IP地址输入、连接/断开、开始/停止监视、导出CSV）。"
        "系统自动识别Modbus TCP和Siemens S7两种协议，连接成功后启动定时采集（默认1000ms间隔）。"
    ),
    make_body(
        "页面中部为实时变量表格，支持在线编辑变量名称、西门子地址、寄存器地址、数据类型、单位、"
        "显示范围等属性，并可通过显示开关控制各通道曲线的显隐。"
        "页面下部为三张趋势曲线图（压力MPa、温度C、流量Nml/min），变量按曲线通道属性自动分组到对应图表。"
        "图表支持左键拖拽平移、滚轮缩放Y轴、鼠标悬停显示Tracker浮层。"
        "通过显示时长设置和自动跟随开关，可灵活控制视口范围。"
        "Y轴始终按当前可见窗口内的数据自适应范围，历史尖峰滑出窗口后Y轴自动缩小。"
        "监视过程中连续3次读取失败会触发自动重连。"
        "停止监视时系统自动计算最终泄漏率并判定合格/不合格。"
    ),
    make_img_placeholder("【图片占位：实时监视界面截图（含趋势曲线和变量表格）】"),
    make_caption("图 7 实时监视界面"),
]

for elem in rm_elements:
    parent.insert(list(parent).index(data_elem2), elem)
print("已插入：实时监视界面章节")

# === 更新原数据分析界面的图编号 ===
for p in doc.paragraphs:
    if p.text.strip() == "图 8 数据分析界面":
        pass  # 已经是图8了
    elif p.text.strip() == "图 6 数据分析界面":
        for run in p.runs:
            run.text = ""
        p.runs[0].text = "图 8 数据分析界面"
        print("数据分析图编号更新为图8")

doc.save("doc/数据分析-管理软件设计方案-A(2).docx")
print("\n文档已保存!")
