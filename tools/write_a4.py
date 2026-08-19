# -*- coding: utf-8 -*-
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

doc = Document('doc/数据分析-管理软件设计方案-A(1).docx')

def norm(t): return t.replace('\xa0', ' ').strip()
def set_text(para, text):
    for r in para.runs: r.text = ''
    if para.runs: para.runs[0].text = text
def make_para(text, tmpl):
    new_p = deepcopy(tmpl._element)
    runs = new_p.findall(qn('w:r'))
    first = True
    for r in runs:
        for t in r.findall(qn('w:t')):
            t.text = text if first else ''
            first = False
    return new_p
def remove_para(p):
    p._element.getparent().remove(p._element)

# 模板
h3_tmpl = list_tmpl = normal_tmpl = caption_tmpl = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 3' and not h3_tmpl: h3_tmpl = p
    if 'List' in (p.style.name or '') and p.runs and not list_tmpl: list_tmpl = p
    if p.style.name == 'Normal' and p.runs and not normal_tmpl: normal_tmpl = p
    if p.style.name == 'Caption' and not caption_tmpl: caption_tmpl = p

# 文本修改
fixes = {
    43: '实现试验数据导入（单文件与批量）、自动归档、自动分类存储；',
    51: '通信层负责实现RS232、RS485、TCP/IP、Modbus TCP、Siemens S7等协议通信；',
    53: '应用层负责向用户提供桌面管理界面、数据分析界面、报表界面以及系统配置界面。',
    58: '首页概览：系统全局数据总览仪表盘，展示核心KPI指标、试验记录预览和系统运行状态。',
    59: '试验对象：管理项目/机组、试验对象路径树（四级层级）、测量装置台账及报告导出。',
    60: '试验路径：管理试验配方，定义泄漏率限值、预充压压力、阀门规格等试验参数。',
    61: '试验记录：存储、查询和管理所有历史试验记录，支持曲线回放与批量操作。',
    62: '实时监视：连接PLC实时采集压力、温度、流量数据，以趋势曲线形式展示。',
    63: '数据分析：提供故障趋势、合格率统计、泄漏率趋势等多维度统计分析。',
    65: '服务器配置：配置不低于4核CPU / 8GB内存 / 256GB SSD，操作系统Windows 10/11或Server 2016+。',
    78: '该模块实现软件与多台测量装置的双向数据交互，支持单文件与批量导入：',
    79: '试验任务下发：将选定的试验对象（含配置参数）下发至指定测量装置。',
    80: '试验数据上传：支持单文件导入（.json/.txt/.csv）和批量导入（选择文件夹自动解析），自动关联试验路径配方。',
    81: '非覆盖存储：同一对象的不同次试验按时间追加，每次导入保存配方快照（JSON）。',
    83: '提供五个分析维度：故障趋势（按阀门类型统计）、合格率统计（各阀门评定）、泄漏率趋势（多系列曲线）、阀门试验次数（Top 50）、机组合格情况（各机组对比）。支持多条件筛选和Excel导出。',
    84: '故障趋势分析：', 85: '按阀门类型统计合格/不合格数，堆叠柱状图展示。',
    86: '合格率统计：展示总试验数、合格率，计算各阀门合格率等级。',
    87: '泄漏率趋势分析：', 88: '多系列趋势曲线展示不同阀门类型的泄漏率历史。',
    89: '阀门试验次数统计：排名展示各阀门试验次数。',
    90: '机组合格情况统计：按机组展示合格率，支持Excel导出。',
    94: '认证：用户名+密码登录，密码加密存储，支持会话管理。',
    95: '审计：所有关键操作记录日志，支持筛选、清理与导出。',
    96: '数据隔离：通过角色权限控制可见页面和可执行操作。',
    98: '通过自动/手动备份、主从自动故障切换确保数据安全与业务连续性。',
    99: '自动备份：可配置间隔定时全量备份，可配置保留天数。',
    101: '主从切换：主库故障自动切至从库，恢复后自动切回，磁盘缓冲确保零丢失。',
    113: '文件操作：导入（JSON/CSV/TXT数据包）、导出（报告/原始数据）。',
    114: '数据对接：封装通信协议（Modbus TCP、Siemens S7、RS232/RS485、USB），提供缓存与断点续传。',
    127: '首页采用"左侧导航+顶部指标+中部监控+底部台账"布局，主要用于：',
    128: '试验数据管理总览', 129: '设备状态监控', 130: '试验合格率统计',
    133: '首页查看6个核心指标（试验对象数、装置数、记录数、合格率、异常数、备份时间）及最近导入详情。',
    134: '采用"左侧导航+顶部指标+中部监控+底部台账"布局结构。',
    136: '关键指标一屏总览；', 137: '试验记录预览与最近导入详情；',
    138: '台账概况（项目/机组/系统/贯穿件/阀门/部件统计）；',
    139: '系统维护状态（数据库连接/备份/同步状态）；',
    140: '核心KPI卡片（试验对象数、装置数、记录数、合格率、异常数、备份状态）；',
    141: '数据分析快速入口。',
    146: '试验对象页面实现基础数据管理，采用四Tab布局：项目/机组、试验对象路径树、测量装置、报告导出。',
    148: '项目/机组信息维护（支持批量CSV导入）；',
    150: '试验对象路径树管理（四级层级：系统→贯穿件→阀门/其他部件）；',
    151: '测量装置台账管理（按通信方式/启用状态筛选）；',
    152: '试验报告导出（Excel/PDF，4种范围）；',
    153: '试验对象与试验路径配方关联管理。',
    154: '四Tab布局使各类基础数据可在同一页面内快速切换维护。',
    157: '测量装置Tab支持查看装置信息和管理操作。报告导出Tab支持Excel和PDF。',
}
for i, t in fixes.items():
    p = doc.paragraphs[i]
    if i < len(doc.paragraphs):
        set_text(p, t)
print(f'文本修改：{len(fixes)} 段')

# 找到界面章节并替换
sec_start = sec_end = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and '概念图详解' in p.text: sec_start = i + 1
    if p.style.name == 'Heading 1' and '系统联调' in p.text: sec_end = i; break
for j in range(sec_end - 1, sec_start - 1, -1):
    remove_para(doc.paragraphs[j])
print(f'删除旧界面章节：{sec_end - sec_start} 段')

# 插入点
target = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and '系统联调' in p.text:
        target = p._element; break
body = target.getparent()

# 新章节 - 按正确顺序
sections = [
    ('h3', '首页概览'),
    ('p', '首页概览采用"左侧导航+顶部指标+中部监控+底部台账"布局，以6个KPI卡片（试验对象数、装置数、记录数、合格率、异常数、备份时间）、试验记录预览、最近导入详情、9项台账统计和系统维护状态为主体，满足全局数据一屏总览。'),
    ('img', '【图片占位：首页概览界面截图】'),
    ('cap', '图 1 首页概览界面'),

    ('h3', '试验对象页面'),
    ('p', '试验对象页面是基础数据管理核心，采用四Tab布局。'),
    ('p', 'Tab 1 - 项目/机组：系统最顶层组织单元。项目归类组织试验数据，机组隶属于项目。支持增删改查（编码自动生成P{年月}{序号}），支持CSV批量导入。'),
    ('p', 'Tab 2 - 试验对象管理：四级层级树（系统→贯穿件→阀门/其他部件）。左侧树形导航，底部快速新建按钮，右侧节点详情及操作，下方试验统计和关联配方。叶子节点可配置默认试验路径。'),
    ('p', 'Tab 3 - 测量装置：表格展示装置信息（编号、名称、IP、通信方式、启用状态）。支持筛选和增删改。仅启用装置可在实时监视中选择。'),
    ('p', 'Tab 4 - 报告导出：支持全部/本月/本月合格/本月不合格四种范围，Excel和PDF格式。'),
    ('img', '【图片占位：试验对象管理界面截图（4个Tab）】'),
    ('cap', '图 2 试验对象管理界面'),

    ('h3', '试验路径页面'),
    ('p', '管理试验配方，定义泄漏率限值、预充压压力、阀门规格等参数。支持搜索、启用过滤、CSV导入导出、增删改。每次编辑自动创建版本快照（RecipeVersion），试验记录保存导入时的配方快照（JSON），后续修改不影响历史。'),
    ('img', '【图片占位：试验路径管理界面截图】'),
    ('cap', '图 3 试验路径管理界面'),

    ('h3', '试验记录页面'),
    ('p', '试验记录页面是核心业务模块，用于试验全过程数据的集中管理、查询、分析与追溯。'),
    ('p', '查询区域支持按项目（级联机组）、结果、时间、关键字组合筛选。分页表格展示：记录编号、项目、机组、对象编码、节点名称、最终泄漏率、泄漏限值、判定结果、试验路径、装置、人员、时间、备注。'),
    ('p', '过程曲线回放：选中记录后加载三张趋势图（压力MPa、温度℃、流量Nml/min），按通道自动分组。支持拖拽平移、滚轮缩放、悬停Tracker、时间范围裁剪。右侧显示通道图例和配方参数。'),
    ('p', '支持批量修改试验路径（自动重算判定）、批量删除、单条编辑、单文件/批量数据导入、多格式报告导出。'),
    ('img', '【图片占位：试验记录界面截图（含曲线回放）】'),
    ('cap', '图 4 试验记录界面'),

    ('h3', '实时监视界面'),
    ('p', '实时监视界面是核心功能模块，连接PLC实时采集压力、温度、流量数据并以趋势曲线展示。'),
    ('p', '顶部控制区：四级级联选择（项目→机组→对象→装置）和PLC连接控制。自动识别Modbus TCP和Siemens S7协议。连续3次读取失败自动重连。'),
    ('p', '中部变量表格：支持在线编辑（名称、地址、类型、单位、范围），显示开关控制曲线显隐。当前值/更新时间/状态只读。修改后保存配置到数据库。'),
    ('p', '底部三张趋势图：压力(MPa)、温度(℃)、流量(Nml/min)，按通道自动分组。支持拖拽/缩放/悬停/显示时长/自动跟随。Y轴按可见窗口自适应。停止时自动计算最终泄漏率并判定合格/不合格。'),
    ('p', '数据安全：周期自动保存（间隔10s→30s→60s→5min动态调整），内存缓冲86400点上限，停止/关闭时同步保存。'),
    ('img', '【图片占位：实时监视界面截图（含趋势曲线和变量表格）】'),
    ('cap', '图 5 实时监视界面'),

    ('h3', '数据分析界面'),
    ('p', '提供五个维度：故障趋势（按阀门类型统计）、合格率统计（各阀门等级评定）、泄漏率趋势（多系列曲线）、阀门试验次数（Top 50排名）、机组合格情况（各机组对比）。支持多条件筛选和Excel导出。'),
    ('img', '【图片占位：数据分析界面截图（5个Tab）】'),
    ('cap', '图 6 数据分析界面'),

    ('h3', '系统设置页面'),
    ('p', '包含五个Tab：用户权限（增删改查、角色分配、启禁用）、角色管理（管理员/工程师/只读三种角色）、操作日志（筛选/清理/导出/保留天数）、数据备份（手动/自动备份、还原、保留策略）、数据库高可用（主从状态、切换参数、连接配置）。'),
    ('img', '【图片占位：系统设置界面截图（5个Tab）】'),
    ('cap', '图 7 系统设置界面'),
]

# 按正确顺序插入
for stype, text in sections:
    tmpl = {'h3': h3_tmpl, 'p': normal_tmpl, 'list': list_tmpl, 'img': normal_tmpl, 'cap': caption_tmpl}[stype]
    elem = make_para(text, tmpl)
    idx = list(body).index(target)
    body.insert(idx, elem)
print(f'插入新章节：{len(sections)} 段')

# 权限表
for table in doc.tables:
    rows = table.rows
    if len(rows) >= 4 and norm(rows[0].cells[0].text) == '角色':
        data = [
            ('只读用户(viewer)', '只读查看与导出', '浏览试验记录、数据分析、导出报告'),
            ('试验工程师(operator)', '试验执行与数据管理', '编辑台账/配方/变量、上传导出数据'),
            ('管理员(admin)', '全权管理', '用户/角色管理、装置管理、备份/恢复、删除数据'),
        ]
        for ri, (c0, c1, c2) in enumerate(data, start=1):
            for ci, val in enumerate([c0, c1, c2]):
                for para in rows[ri].cells[ci].paragraphs:
                    for r in para.runs: r.text = ''
                    if para.runs: para.runs[0].text = val
                    break
        print('权限表已更新')
        break

# 图片占位
for i, p in enumerate(doc.paragraphs):
    drawings = p._element.findall('.//' + qn('w:drawing'))
    if not drawings: continue
    next_t = norm(doc.paragraphs[i+1].text) if i+1 < len(doc.paragraphs) else ''
    if '图 1' in next_t: ph = '【图片占位：首页概览界面截图】'
    elif '图 2' in next_t: ph = '【图片占位：试验对象管理界面截图】'
    elif '图 3' in next_t: ph = '【图片占位：试验路径管理界面截图】'
    elif i < 15: ph = '【图片占位：封面Logo/装饰图】'
    else: ph = '【图片占位：软件主界面概念图】'
    for r in p.runs: r.text = ''
    for r_elem in p._element.findall(qn('w:r')):
        if r_elem.findall('.//' + qn('w:drawing')): p._element.remove(r_elem)
    if p.runs: p.runs[0].text = ph
    else: p.add_run(ph)

# 保存
OUTPUT = 'doc/数据分析-管理软件设计方案-A(4).docx'
tmp = OUTPUT + '.tmp'
doc.save(tmp)
try: os.remove(OUTPUT)
except: pass
os.rename(tmp, OUTPUT)
print(f'\n已保存到：{OUTPUT} ({os.path.getsize(OUTPUT)/1024:.1f} KB)')
