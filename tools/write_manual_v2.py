# -*- coding: utf-8 -*-
"""
重写用户操作手册 - 操作导向版本
重点：步骤化、动作化、用户视角
"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============ 页面和样式设置（同前）============
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(3)
rPr = style._element.find(qn('w:rPr'))
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), '宋体')
rPr.append(rFonts)

for lvl in [1, 2, 3]:
    h = doc.styles[f'Heading {lvl}']
    h.font.name = '宋体'
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    sizes = {1: 18, 2: 16, 3: 15}
    h.font.size = Pt(sizes[lvl])
    if lvl == 1:
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rPr = h._element.find(qn('w:rPr'))
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)

# ============ 工具函数 ============
def add_step(num, text):
    """添加操作步骤"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(f'{num}. {text}')
    run.font.size = Pt(12)

def add_body(text):
    """正文段落"""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            clean = part.replace('`', '').strip()
            if clean:
                p.add_run(clean)

def add_note(text):
    """注意事项"""
    p = doc.add_paragraph()
    run = p.add_run(f'⚠ {text}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

def add_result(text):
    """操作结果"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(f'→ {text}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0x66, 0x00)

def add_image_placeholder(desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'【图片占位：{desc}】')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text).font.size = Pt(11)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            cell.paragraphs[0].add_run(val).font.size = Pt(11)

# ============ 正文内容 ============

# 封面
doc.add_heading('智能安全壳隔离阀泄漏率数据管理软件', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('用户操作手册')
run.font.size = Pt(16)
run.bold = True

doc.add_paragraph()
add_body('软件版本：B 版')
add_body('文档日期：2026 年 7 月')
doc.add_paragraph()

# ==================== 目录 ====================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('目  录')
run.font.size = Pt(18)
run.bold = True
doc.add_paragraph()

# 分页
doc.add_page_break()

# ==================== 1. 快速入门 ====================
doc.add_heading('1. 快速入门', level=2)

doc.add_heading('1.1 启动与登录', level=3)
add_step(1, '双击桌面图标或从开始菜单启动软件')
add_step(2, '在登录界面输入用户名和密码')
add_step(3, '点击"登录"按钮')
add_result('登录成功，进入首页概览页面')
add_note('连续5次输入错误密码，账户将被锁定30分钟')

add_image_placeholder('登录界面截图')
add_caption('图 1 登录界面')

doc.add_heading('1.2 软件界面说明', level=3)
add_body('软件界面分为左右两部分：')
add_step(1, '左侧导航栏：包含7个功能页面入口，点击切换页面')
add_step(2, '右侧内容区：显示当前选中页面的功能界面')
add_body('底部状态栏显示：数据库连接状态（主库/从库）、当前登录用户信息。')

add_image_placeholder('软件主界面截图')
add_caption('图 2 软件主界面')

doc.add_heading('1.3 首次使用流程', level=3)
add_body('首次使用软件，请按以下顺序完成基础配置：')
add_step(1, '进入"试验对象"页面，在"项目/机组"Tab中新增项目和机组')
add_step(2, '切换到"试验对象管理"Tab，创建试验对象路径树（系统→贯穿件→阀门）')
add_step(3, '切换到"测量装置"Tab，新增测量装置并设置启用状态')
add_step(4, '进入"试验路径"页面，配置试验配方（泄漏率限值、压力参数等）')
add_step(5, '进入"实时监视"页面，选择试验对象和装置，连接PLC，开始监视')

# ==================== 2. 试验对象管理 ====================
doc.add_heading('2. 试验对象管理', level=2)

doc.add_heading('2.1 如何新增项目', level=3)
add_step(1, '点击左侧导航栏"试验对象"')
add_step(2, '确认当前在"项目/机组"Tab（如不在，点击切换）')
add_step(3, '点击左上角"新增项目"按钮')
add_step(4, '在弹出窗口中输入项目名称')
add_step(5, '选择启用状态（启用/停用）')
add_step(6, '输入备注信息（可选）')
add_step(7, '点击"确定"保存')
add_result('项目列表中出现新项目，编码自动生成（格式：P+年月+序号）')

add_image_placeholder('新增项目对话框截图')
add_caption('图 3 新增项目对话框')

doc.add_heading('2.2 如何新增机组', level=3)
add_step(1, '在项目列表中点击选中目标项目')
add_step(2, '右侧自动显示该项目下的机组列表')
add_step(3, '点击右上角"新增机组"按钮')
add_step(4, '输入机组名称，选择状态，输入备注')
add_step(5, '点击"确定"保存')
add_result('机组列表中出现新机组，自动关联到选中的项目')

doc.add_heading('2.3 如何批量导入项目/机组', level=3)
add_body('使用CSV文件批量导入多个项目和机组：')
add_step(1, '在"项目/机组"Tab中，点击"批量导入"按钮')
add_step(2, '选择包含CSV文件的文件夹')
add_step(3, '系统自动解析文件，显示导入进度条')
add_step(4, '导入完成后，查看导入结果（成功/失败条数）')
add_note('CSV文件格式要求：第一行为表头，包含"项目编码"、"项目名称"、"机组编码"、"机组名称"等列')

doc.add_heading('2.4 如何创建试验对象路径树', level=3)
add_step(1, '切换到"试验对象管理"Tab')
add_step(2, '在顶部下拉框中选择项目和机组（限定范围）')
add_step(3, '点击底部"新建系统"按钮，输入系统名称，点击确定')
add_step(4, '在树形列表中选中刚创建的系统节点')
add_step(5, '点击"新建贯穿件"按钮，输入名称，点击确定')
add_step(6, '选中贯穿件节点，点击"新建阀门"或"新建其他部件"')
add_step(7, '在弹出窗口中配置：节点名称、泄漏率限值、试验压力、默认关联试验路径')
add_step(8, '点击"确定"保存')
add_result('路径树中显示完整的四级层级结构')

add_image_placeholder('试验对象路径树界面截图')
add_caption('图 4 试验对象路径树')

add_note('只有叶子节点（阀门/其他部件）才能进行试验。泄漏率限值用于自动判定合格/不合格。')

doc.add_heading('2.5 如何新增测量装置', level=3)
add_step(1, '切换到"测量装置"Tab')
add_step(2, '点击"新增装置"按钮')
add_step(3, '填写装置信息：')
add_body('  - 装置名称（必填）')
add_body('  - IP地址（如使用网络通信）')
add_body('  - 通信方式（USB/RJ45/RS232/RS485）')
add_body('  - 启用状态（选择"启用"才能在实时监视中使用）')
add_step(4, '点击"确定"保存')
add_result('装置列表中出现新装置，编号自动生成（格式：DEV+时间戳）')

add_note('实时监视中选择装置时，只有"启用"状态的装置才会出现在下拉列表中。')

doc.add_heading('2.6 如何导出试验报告', level=3)
add_step(1, '切换到"报告导出"Tab')
add_step(2, '在"导出范围"下拉框中选择：全部/本月/本月合格/本月不合格')
add_step(3, '在"导出格式"下拉框中选择：Excel 或 PDF')
add_step(4, '选择导出目录（点击"浏览"按钮）')
add_step(5, '点击"导出"按钮')
add_step(6, '等待导出完成，查看导出结果')
add_result('指定目录下生成报告文件，包含试验记录和相关数据')

add_image_placeholder('报告导出界面截图')
add_caption('图 5 报告导出界面')

# ==================== 3. 试验路径管理 ====================
doc.add_heading('3. 试验路径（配方）管理', level=2)

doc.add_heading('3.1 如何新增试验路径', level=3)
add_step(1, '点击左侧导航栏"试验路径"')
add_step(2, '点击右上角"新增"按钮')
add_step(3, '在弹出窗口中填写：')
add_body('  - 试验路径名称（必填，唯一）')
add_body('  - 所属系统')
add_body('  - 贯穿件直径（mm）')
add_body('  - 试验阀门编号')
add_body('  - 阀门公称直径（mm）')
add_body('  - 泄漏率设计最大值（Nml/min）')
add_body('  - 预充压压力P2（MPa）')
add_step(4, '选择"启用"状态')
add_step(5, '输入备注（可选）')
add_step(6, '点击"保存"')
add_result('试验路径列表中出现新配方，自动创建版本快照（V1）')

add_image_placeholder('试验路径编辑界面截图')
add_caption('图 6 试验路径编辑界面')

add_note('每次修改配方并保存后，系统自动创建新版本快照。历史试验记录保存的是导入时的配方快照，不受后续修改影响。')

doc.add_heading('3.2 如何通过CSV批量导入配方', level=3)
add_step(1, '在"试验路径"页面，点击"导入CSV"按钮')
add_step(2, '选择CSV文件（支持GBK或UTF-8编码）')
add_step(3, '系统自动解析并导入，显示导入结果')
add_result('导入成功条数、失败条数、跳过条数显示在界面上')

doc.add_heading('3.3 如何导出配方列表', level=3)
add_step(1, '点击"导出CSV"按钮')
add_step(2, '选择保存位置和文件名')
add_step(3, '点击"保存"')
add_result('配方列表导出为CSV文件，包含所有字段')

# ==================== 4. 试验记录管理 ====================
doc.add_heading('4. 试验记录管理', level=2)

doc.add_heading('4.1 如何查询试验记录', level=3)
add_step(1, '点击左侧导航栏"试验记录"')
add_step(2, '在顶部查询区域设置筛选条件：')
add_body('  - 选择项目（可选）')
add_body('  - 选择机组（可选，按项目级联过滤）')
add_body('  - 选择结果（全部/合格/不合格/未知）')
add_body('  - 设置日期范围（起始日期和结束日期）')
add_body('  - 输入关键字（记录编号/对象名称/装置编号）')
add_step(3, '点击"查询"按钮')
add_result('下方表格显示符合条件的记录，分页展示')

add_image_placeholder('试验记录查询界面截图')
add_caption('图 7 试验记录查询界面')

doc.add_heading('4.2 如何查看过程曲线', level=3)
add_step(1, '在试验记录列表中，点击选中一条记录')
add_step(2, '页面下方自动加载该记录的三张趋势曲线图（压力/温度/流量）')
add_step(3, '鼠标悬停在曲线上，查看具体数值')
add_step(4, '使用鼠标滚轮缩放曲线')
add_step(5, '按住鼠标左键拖拽，平移时间轴')
add_step(6, '在时间范围输入框中输入起止时间（秒），点击"确认"裁剪显示范围')
add_result('右侧面板显示通道图例和配方参数')

add_note('输入时间范围时，0表示从头开始，0表示到末尾。例如"0-3600"表示显示前1小时的数据。')

doc.add_heading('4.3 如何导入试验数据（单文件）', level=3)
add_step(1, '点击"导入数据"按钮')
add_step(2, '在文件选择器中选择数据包文件（.json/.txt/.csv）')
add_step(3, '系统自动解析文件，识别试验对象')
add_step(4, '确认或修改：项目编码、机组编码、操作人员')
add_step(5, '确认或选择试验路径（系统自动匹配默认配方）')
add_step(6, '如果文件不含装置编号，在弹出窗口中选择测量装置')
add_step(7, '点击"开始上传"')
add_result('系统创建试验记录，自动判定合格/不合格，状态显示为"导入成功"')

add_note('如果试验对象已配置默认试验路径，系统会自动关联该配方。否则需要手动选择。')

doc.add_heading('4.4 如何批量导入试验数据', level=3)
add_step(1, '点击"批量导入"按钮')
add_step(2, '选择包含多个数据包文件的文件夹')
add_step(3, '系统自动解析所有文件，显示解析进度')
add_step(4, '查看解析结果列表：')
add_body('  - 已匹配：显示对应的项目、机组、试验对象')
add_body('  - 未匹配：显示"未匹配"，需要手动指定')
add_step(5, '对于未匹配项，手动选择对应的项目/机组/对象')
add_step(6, '点击"批量上传"')
add_step(7, '查看批量导入进度和结果')
add_result('导入完成的记录显示在试验记录列表中')

doc.add_heading('4.5 如何批量修改试验路径', level=3)
add_step(1, '在查询结果中，勾选多条记录（或点击表头全选）')
add_step(2, '点击"批量修改试验路径"按钮')
add_step(3, '在弹出窗口中选择新的试验路径')
add_step(4, '点击"确定"')
add_result('所有选中记录的试验路径已更新，系统自动重新计算合格/不合格判定')

add_note('修改试验路径后，系统会根据新配方的泄漏率限值重新判定。如果原合格变为不合格，状态会更新。')

doc.add_heading('4.6 如何删除试验记录', level=3)
add_body('删除单条记录：')
add_step(1, '选中要删除的记录')
add_step(2, '点击"删除"按钮')
add_step(3, '在确认对话框中点击"确定"')
add_result('记录从列表中删除')

add_body('批量删除：')
add_step(1, '勾选多条记录')
add_step(2, '点击"批量删除"按钮')
add_step(3, '在确认对话框中点击"确定"')
add_result('所有选中记录被删除')

add_note('删除操作不可恢复，请谨慎操作。系统会弹出二次确认对话框。')

# ==================== 5. 实时监视 ====================
doc.add_heading('5. 实时监视', level=2)

doc.add_heading('5.1 如何开始实时监视', level=3)
add_step(1, '点击左侧导航栏"实时监视"')
add_step(2, '在顶部控制区，依次选择：')
add_body('  - 项目（下拉框）')
add_body('  - 机组（按项目过滤）')
add_body('  - 试验对象（按机组过滤，显示编码+名称）')
add_body('  - 测量装置（仅显示启用状态的装置）')
add_step(3, '在PLC地址输入框中输入PLC的IP地址（默认127.0.0.1）')
add_step(4, '点击"保存地址"（可选，保存后下次启动自动使用）')
add_step(5, '点击"连接PLC"')
add_step(6, '等待连接成功（状态栏显示"已连接"，指示灯变绿）')
add_step(7, '点击"开始监视"')
add_result('系统创建试验记录，开始实时采集数据，曲线图开始绘制')

add_image_placeholder('实时监视界面截图')
add_caption('图 8 实时监视界面')

add_note('如果连接失败，检查PLC的IP地址和网络连通性。如果启用了仿真模式，会自动使用模拟数据。')

doc.add_heading('5.2 如何添加/编辑监控变量', level=3)
add_step(1, '在实时监视页面，点击"变量配置"按钮')
add_step(2, '在弹出窗口中点击"新增变量"')
add_step(3, '填写变量信息：')
add_body('  - 变量名称')
add_body('  - 西门子地址（如 DB15.DBD0）或 寄存器地址（如 512）')
add_body('  - 数据类型（double/int/float/real/ushort/dword）')
add_body('  - 单位（MPa/℃/Nml/min 等）')
add_body('  - 曲线分组（压力/温度/流量）')
add_body('  - 最小值/最大值（Y轴显示范围）')
add_step(4, '点击"保存"')
add_result('变量列表中出现新变量，曲线图中显示对应曲线')

add_note('修改变量配置后，必须点击"保存配置"才能生效。重启软件后配置会丢失。')

doc.add_heading('5.3 如何控制曲线显示', level=3)
add_body('显示/隐藏曲线：')
add_step(1, '在变量表格中，点击"显示"列的复选框')
add_step(2, '勾选=显示曲线，取消勾选=隐藏曲线')

add_body('调整显示时长：')
add_step(1, '在"显示时长"输入框中输入秒数（默认600秒=10分钟）')
add_step(2, '点击"确认"')
add_result('X轴显示范围更新为最近N秒的数据')

add_body('启用/禁用自动跟随：')
add_step(1, '点击"自动"复选框')
add_step(2, '勾选=视口跟随最新数据滚动，Y轴按可见窗口自适应')
add_step(3, '取消勾选=视口停在当前位置，可手动拖拽查看历史数据')

doc.add_heading('5.4 如何停止监视', level=3)
add_step(1, '点击"停止监视"按钮')
add_step(2, '在确认对话框中点击"确定"')
add_result('系统保存所有采集数据，计算最终泄漏率，判定合格/不合格，更新试验记录状态')

add_note('停止监视后，可以在"试验记录"页面查看该次试验的详细数据和曲线。')

doc.add_heading('5.5 如何导出实时数据', level=3)
add_step(1, '在实时监视页面，点击"导出CSV"按钮')
add_step(2, '选择保存位置和文件名')
add_step(3, '点击"保存"')
add_result('所有采集的数据导出为CSV文件，包含时间戳和所有变量值')

doc.add_heading('5.6 自动重连机制', level=3)
add_body('监视过程中，如果PLC连接中断：')
add_step(1, '系统自动尝试重连（最多3次）')
add_step(2, '如果重连成功，继续监视，状态栏显示"已重连"')
add_step(3, '如果3次重连都失败，停止监视，状态栏显示"连接失败"')

add_note('重连失败后，需要手动检查PLC状态，重新连接后再开始监视。')

# ==================== 6. 数据分析 ====================
doc.add_heading('6. 数据分析', level=2)

doc.add_heading('6.1 如何进行数据分析', level=3)
add_step(1, '点击左侧导航栏"数据分析"')
add_step(2, '在顶部筛选区域设置条件：')
add_body('  - 选择项目（可选）')
add_body('  - 选择机组（可选）')
add_body('  - 选择系统（可选）')
add_body('  - 设置时间范围（起始日期和结束日期）')
add_step(3, '点击"查询"按钮')
add_step(4, '查看5个分析维度的结果：')
add_body('  - Tab 1 故障趋势：查看按阀门类型统计的合格/不合格数')
add_body('  - Tab 2 合格率统计：查看各阀门的合格率')
add_body('  - Tab 3 泄漏率趋势：查看不同阀门类型的泄漏率变化曲线')
add_body('  - Tab 4 阀门试验次数：查看试验次数排名')
add_body('  - Tab 5 机组合格情况：查看各机组的合格率对比')
add_step(5, '点击"导出Excel"按钮，保存分析结果')
add_result('生成包含多个Sheet的Excel文件，包含所有分析数据')

add_image_placeholder('数据分析界面截图')
add_caption('图 9 数据分析界面')

# ==================== 7. 系统设置 ====================
doc.add_heading('7. 系统设置', level=2)

doc.add_heading('7.1 如何新增用户', level=3)
add_step(1, '点击左侧导航栏"系统设置"')
add_step(2, '切换到"用户管理"Tab')
add_step(3, '点击"新增用户"按钮')
add_step(4, '填写：用户名、密码、姓名')
add_step(5, '选择角色（可多选：管理员/试验工程师/只读用户）')
add_step(6, '选择部门')
add_step(7, '点击"保存"')
add_result('用户列表中出现新用户')

add_note('只有管理员角色才能访问"系统设置"页面。')

doc.add_heading('7.2 如何手动备份数据库', level=3)
add_step(1, '切换到"数据备份"Tab')
add_step(2, '点击"立即备份"按钮')
add_step(3, '等待备份完成（状态栏显示进度）')
add_result('备份文件保存在配置的备份目录中，文件名包含时间戳')

doc.add_heading('7.3 如何配置自动备份', level=3)
add_step(1, '在"数据备份"Tab中，点击"备份设置"')
add_step(2, '勾选"启用自动备份"')
add_step(3, '设置备份间隔（小时）')
add_step(4, '设置保留天数（超过天数的备份自动删除）')
add_step(5, '点击"保存"')
add_result('系统按配置间隔自动备份数据库')

doc.add_heading('7.4 如何配置数据库高可用', level=3)
add_step(1, '切换到"数据库高可用"Tab')
add_step(2, '点击"配置"按钮')
add_step(3, '填写从库信息：')
add_body('  - 服务器地址（IP或主机名）')
add_body('  - 端口（默认1433）')
add_body('  - 用户名和密码')
add_step(4, '点击"测试连接"验证从库可访问')
add_step(5, '配置故障切换参数：')
add_body('  - 健康检测间隔（默认15秒）')
add_body('  - 失败重试次数（默认3次）')
add_body('  - 切换超时时间（默认30秒）')
add_step(6, '点击"保存"')
add_result('主从切换配置生效，主库故障时自动切换到从库')

add_note('配置高可用前，请确保从库已配置并正常运行。')

# ==================== 8. 常见问题 ====================
doc.add_heading('8. 常见问题与故障排除', level=2)

doc.add_heading('8.1 PLC连接失败', level=3)
add_body('问题：点击"连接PLC"后，状态栏显示"连接失败"')
add_body('可能原因：')
add_body('  - PLC的IP地址不正确')
add_body('  - 网络不通（检查网线、交换机、防火墙）')
add_body('  - PLC未上电或未启动')
add_body('  - PLC的通信端口未开放')
add_body('解决方法：')
add_step(1, '检查PLC的IP地址是否正确')
add_step(2, '使用ping命令测试网络连通性')
add_step(3, '检查PLC电源和运行状态')
add_step(4, '确认PLC的通信端口（Modbus TCP默认502，S7默认102）已开放')
add_step(5, '检查防火墙设置，确保端口未被阻止')

doc.add_heading('8.2 监视过程中数据不更新', level=3)
add_body('问题：开始监视后，曲线图数据不更新')
add_body('可能原因：')
add_body('  - PLC通信中断')
add_body('  - 采样间隔设置过长')
add_body('  - 变量地址配置错误')
add_body('解决方法：')
add_step(1, '检查PLC连接状态，如断开则重新连接')
add_step(2, '检查采样间隔设置（建议1000ms）')
add_step(3, '检查变量配置中的地址是否正确')
add_step(4, '停止监视，重新开始')

doc.add_heading('8.3 数据库连接失败', level=3)
add_body('问题：启动软件或操作时提示"数据库连接失败"')
add_body('可能原因：')
add_body('  - SQL Server服务未启动')
add_body('  - 数据库服务器地址配置错误')
add_body('  - 用户名或密码错误')
add_body('  - 网络问题')
add_body('解决方法：')
add_step(1, '检查SQL Server服务是否启动（Windows服务管理器）')
add_step(2, '检查配置文件中的数据库连接字符串')
add_step(3, '使用SQL Server Management Studio测试连接')
add_step(4, '如果配置了主从切换，检查从库是否可用')

doc.add_heading('8.4 数据导入失败', level=3)
add_body('问题：导入试验数据时提示失败')
add_body('可能原因：')
add_body('  - 文件格式不正确')
add_body('  - 试验对象未创建')
add_body('  - 测量装置未配置')
add_body('解决方法：')
add_step(1, '检查文件格式是否符合要求（JSON/TXT/CSV）')
add_step(2, '确认试验对象已在"试验对象管理"中创建')
add_step(3, '确认测量装置已在"测量装置"Tab中配置')
add_step(4, '检查文件内容是否完整，无损坏')

doc.add_heading('8.5 曲线回放无数据', level=3)
add_body('问题：在试验记录中查看曲线时，显示"无数据"')
add_body('可能原因：')
add_body('  - 监视过程中未采集到数据')
add_body('  - 监视时间过短（少于1秒）')
add_body('  - 数据库存储失败')
add_body('解决方法：')
add_step(1, '检查监视过程中是否有数据点')
add_step(2, '检查监视时长是否足够')
add_step(3, '查看系统日志，确认数据是否成功存储')

# ==================== 附录 ====================
doc.add_heading('附录', level=2)

doc.add_heading('A. 快捷键', level=3)
add_table(
    ['快捷键', '功能'],
    [
        ['Ctrl+F', '打开查询对话框'],
        ['F5', '刷新当前页面'],
        ['Ctrl+S', '保存配置'],
        ['Ctrl+E', '导出数据'],
        ['Esc', '关闭弹出窗口'],
    ]
)

doc.add_heading('B. 术语表', level=3)
add_table(
    ['术语', '说明'],
    [
        ['试验路径/配方', '定义试验参数的配置，包括泄漏率限值、压力等'],
        ['试验对象', '被测试的设备或部件，组织成四级层级树'],
        ['贯穿件', '穿过安全壳的管道或电缆通道'],
        ['泄漏率', '衡量密封性能的指标，单位Nml/min'],
        ['预充压P2', '试验前的充压压力，单位MPa'],
        ['主库/从库', '数据库高可用架构中的主数据库和备用数据库'],
        ['磁盘缓冲', '数据库故障时的临时数据存储机制'],
    ]
)

doc.add_heading('C. 技术支持', level=3)
add_body('如遇到本文档未覆盖的问题，请联系技术支持：')
add_body('  - 电话：XXX-XXXX-XXXX')
add_body('  - 邮箱：support@example.com')
add_body('  - 工作时间：周一至周五 9:00-17:00')

# ============ 保存 ============
OUTPUT = 'doc/用户操作手册.docx'
doc.save(OUTPUT)
print(f'已保存到：{OUTPUT}')
print(f'文件大小：{os.path.getsize(OUTPUT)/1024:.1f} KB')
