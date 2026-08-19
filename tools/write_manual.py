# -*- coding: utf-8 -*-
"""
重写用户操作手册Word文档，与设计方案文档格式一致。
字体：正文宋体10.5pt两端对齐，标题黑体粗体居中。
图片位置添加占位符。
"""
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ============ 页面设置 ============
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ============ 样式设置（与设计方案一致）============
# Normal: 宋体 10.5pt 两端对齐
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.font.bold = False
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(3)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Heading 1: 黑体 18pt 粗体
h1 = doc.styles['Heading 1']
h1.font.name = '黑体'
h1.font.size = Pt(18)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(17)
h1.paragraph_format.space_after = Pt(17)
h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Heading 2: 黑体 16pt 粗体
h2 = doc.styles['Heading 2']
h2.font.name = '黑体'
h2.font.size = Pt(16)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(13)
h2.paragraph_format.space_after = Pt(13)
h2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Heading 3: 黑体 15pt 粗体
h3 = doc.styles['Heading 3']
h3.font.name = '黑体'
h3.font.size = Pt(15)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.paragraph_format.line_spacing = 1.73
h3.paragraph_format.space_before = Pt(13)
h3.paragraph_format.space_after = Pt(13)
h3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# ============ 工具函数 ============
def add_body(text):
    """正文段落，处理 **粗体**"""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            clean = part.replace('`', '').replace('⚠️', '').strip()
            if clean:
                p.add_run(clean)
    return p

def add_list(text):
    """列表项"""
    clean = text.replace('**', '').replace('`', '').replace('⚠️', '').strip()
    p = doc.add_paragraph(clean, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_table(headers, rows):
    """表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val.replace('`', ''))
            run.font.size = Pt(10)
            run.font.name = '宋体'

def add_image_placeholder(desc):
    """图片占位符"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'【图片占位：{desc}】')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

def add_caption(text):
    """图标题"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(10.5)

# ============ 写入内容 ============

# 标题
doc.add_heading('智能安全壳隔离阀泄漏率数据管理软件', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('用户操作手册')
run.font.size = Pt(16)
run.bold = True

doc.add_paragraph()
add_body('软件版本：B 版')
add_body('适用设备：智能安全壳隔离阀泄漏率测量装置')
add_body('文档日期：2026 年 7 月')
doc.add_paragraph()

# === 1. 软件概述 ===
doc.add_heading('1. 软件概述', level=2)
doc.add_heading('1.1 软件用途', level=3)
add_body('本软件用于核安全壳隔离阀泄漏率试验的全流程数据管理，涵盖：基础台账维护（项目、机组、试验对象路径树、测量装置）、试验路径（配方）配置与版本管理、现场数据包导入（单文件/批量）、实时采集 PLC 数据并绘制趋势曲线、试验记录管理与判定、多维度历史数据统计分析、试验报告导出（Excel/PDF）、数据库主从高可用与自动备份。')

doc.add_heading('1.2 界面布局', level=3)
add_body('软件采用"左侧导航 + 右侧内容"布局。左侧导航栏包含7个页面入口（首页概览、试验对象、试验路径、试验记录、实时监视、数据分析、系统设置），底部显示数据库主/从库状态和当前登录用户信息。右侧为当前页面的功能界面。软件启动后自动最大化，自适应高分屏缩放。')
add_image_placeholder('软件主界面整体布局截图')
add_caption('图 1 软件主界面')

# === 2. 登录与权限 ===
doc.add_heading('2. 登录与权限', level=2)
doc.add_heading('2.1 用户登录', level=3)
add_body('启动软件后进入登录界面，输入用户名和密码后点击"登录"。')

doc.add_heading('2.2 角色与权限', level=3)
add_body('系统内置三种角色，导航栏按权限过滤，按钮按权限启用/禁用：')
add_table(
    ['功能', '管理员(admin)', '试验工程师(operator)', '只读用户(viewer)'],
    [
        ['首页概览', '✅', '✅', '✅'],
        ['试验对象（查看）', '✅', '✅', '✅'],
        ['试验对象（编辑）', '✅', '✅', '❌'],
        ['试验对象（删除）', '✅', '❌', '❌'],
        ['试验路径（查看/编辑）', '✅', '✅', '❌'],
        ['试验路径（删除）', '✅', '❌', '❌'],
        ['试验记录（查看/导出）', '✅', '✅', '✅'],
        ['试验记录（上传/修改）', '✅', '✅', '❌'],
        ['试验记录（删除）', '✅', '❌', '❌'],
        ['实时监视（查看/编辑）', '✅', '✅', '❌'],
        ['实时监视（删除）', '✅', '❌', '❌'],
        ['数据分析', '✅', '✅', '✅'],
        ['系统设置', '✅', '❌', '❌'],
    ]
)

# === 3. 首页概览 ===
doc.add_heading('3. 首页概览', level=2)
add_body('首页是软件启动后的默认页面，提供全局数据一览。采用"左侧导航 + 顶部指标 + 中部监控 + 底部台账"布局。')

add_image_placeholder('首页概览界面截图')
add_caption('图 2 首页概览界面')

doc.add_heading('3.1 核心指标卡片（顶部6格）', level=3)
add_table(
    ['卡片', '说明'],
    [
        ['试验对象', '系统中阀门 + 部件数量'],
        ['测量装置', '启用状态装置数量'],
        ['历史记录', '全部试验记录条数'],
        ['本月合格率', '最近30天合格比率(%)'],
        ['待处理异常', '最近30天不合格数'],
        ['最近备份', '最近一次备份时间'],
    ]
)

doc.add_heading('3.2 其他区域', level=3)
add_body('中部左侧为试验记录预览表（最近5条），中部右侧为最近导入详情。底部左侧为9项台账统计（项目/机组/系统/贯穿件/阀门/部件/记录/合格/不合格），右侧为系统维护状态（数据库连接、备份状态、同步状态）。')

# === 4. 试验对象 ===
doc.add_heading('4. 试验对象', level=2)
add_body('试验对象页面是基础数据管理核心，采用四Tab布局。')

add_image_placeholder('试验对象管理界面截图（4个Tab）')
add_caption('图 3 试验对象管理界面')

doc.add_heading('4.1 Tab 1：项目/机组', level=3)
add_body('系统最顶层组织单元。左侧为项目列表，右侧为机组列表（按选中项目过滤）。支持增删改查，项目编码自动生成（格式P{年月}{序号}），支持CSV批量导入。')

doc.add_heading('4.2 Tab 2：试验对象管理（路径树）', level=3)
add_body('采用四级层级树结构：系统(System) → 贯穿件(Penetration) → 阀门(Valve)/其他部件(OtherComponent)。左侧树形导航，底部四个快速新建按钮，右侧显示节点详情（编号、名称、类型、泄漏率限值、试验压力、备注）及操作按钮（修改、导入、导出、删除）。下方展示试验统计和关联配方信息。叶子节点可配置默认试验路径。')

doc.add_heading('4.3 Tab 3：测量装置', level=3)
add_body('表格展示所有装置信息（编号、名称、IP、通信方式、启用状态、连接状态、同步/导入时间）。支持按通信方式和启用状态筛选。提供新增、编辑、删除操作。仅启用状态的装置可在实时监视中选择。')

doc.add_heading('4.4 Tab 4：报告导出', level=3)
add_body('支持导出全部/本月/本月合格/本月不合格四种范围，Excel和PDF两种格式。可自定义文件名和导出目录，提供快速导出按钮。')

# === 5. 试验路径 ===
doc.add_heading('5. 试验路径', level=2)
add_body('管理试验配方，定义泄漏率限值、预充压压力、阀门规格等参数。支持搜索、启用过滤、CSV导入导出、增删改。每次编辑自动创建版本快照（RecipeVersion），试验记录保存导入时的配方快照（JSON），后续修改配方不影响历史。')

add_image_placeholder('试验路径管理界面截图')
add_caption('图 4 试验路径管理界面')

add_body('配方编辑对话框分三个区域：基础信息（名称、系统、启用状态、备注）、阀门参数（贯穿件直径、阀门编号、公称直径）、试验参数（泄漏率设计最大值、预充压P2）。')

# === 6. 试验记录 ===
doc.add_heading('6. 试验记录', level=2)
add_body('核心业务模块，用于试验全过程数据的集中管理、查询、分析与追溯。')

add_image_placeholder('试验记录界面截图（含曲线回放）')
add_caption('图 5 试验记录界面')

doc.add_heading('6.1 查询与列表', level=3)
add_body('顶部支持按项目（级联机组）、结果、时间、关键字组合筛选。结果以分页表格展示，含记录编号、项目、机组、对象编码、节点名称、最终泄漏率、泄漏限值、判定结果（合格绿/不合格红）、试验路径、装置、人员、时间、备注等字段。支持表头全选批量选择。')

doc.add_heading('6.2 过程曲线回放', level=3)
add_body('选中记录后底部自动加载三张趋势曲线图（压力MPa、温度℃、流量Nml/min），按通道自动分组。支持拖拽平移、滚轮缩放、悬停Tracker、时间范围裁剪。右侧显示通道图例和配方参数。')

doc.add_heading('6.3 批量操作与数据导入', level=3)
add_body('支持批量修改试验路径（自动重算判定）、批量删除（二次确认）、双击编辑单条记录。支持单文件导入（.json/.txt/.csv，自动识别对象并关联默认配方）和批量导入（选择文件夹自动解析匹配）。导出支持Excel/PDF/CSV。')

# === 7. 实时监视 ===
doc.add_heading('7. 实时监视', level=2)
add_body('核心功能模块，连接PLC实时采集压力、温度、流量数据并以趋势曲线展示。')

add_image_placeholder('实时监视界面截图（含趋势曲线和变量表格）')
add_caption('图 6 实时监视界面')

doc.add_heading('7.1 控制区与PLC连接', level=3)
add_body('顶部四级级联选择（项目→机组→对象→装置）。PLC连接支持Modbus TCP和Siemens S7协议自动识别。提供保存地址、连接/断开、开始/停止监视、导出CSV按钮。连续3次读取失败自动重连。')

doc.add_heading('7.2 变量表格', level=3)
add_body('中部实时变量表格支持在线编辑：显示开关、颜色、变量名称、西门子地址、寄存器地址、数据类型、单位、最小/最大值。当前值/更新时间/状态只读。修改后点"保存配置"持久化到数据库。')

doc.add_heading('7.3 趋势曲线', level=3)
add_body('底部三张图按通道分组：压力(MPa)、温度(℃)、流量(Nml/min)。支持拖拽/缩放/悬停。通过显示时长输入框（默认600秒）和"自动"复选框控制视口。勾选自动时跟随最新数据，Y轴按可见窗口自适应；取消时停在当前位置。所有数据始终保留。')

doc.add_heading('7.4 数据安全', level=3)
add_body('周期自动保存（间隔随数据量调整：10s→30s→60s→5min），内存缓冲86400点上限，停止/关闭时同步保存。停止监视时自动计算最终泄漏率并判定合格/不合格。')

# === 8. 数据分析 ===
doc.add_heading('8. 数据分析', level=2)
add_body('提供五个分析维度：故障趋势（按阀门类型统计合格/不合格，堆叠柱状图）、合格率统计（各阀门合格率等级评定）、泄漏率趋势（多系列曲线对比）、阀门试验次数（Top 50排名+柱状图）、机组合格情况（各机组合格率对比）。支持按项目/机组/系统/时间筛选，可导出多Sheet的Excel文件。')

add_image_placeholder('数据分析界面截图（5个Tab）')
add_caption('图 7 数据分析界面')

# === 9. 系统设置 ===
doc.add_heading('9. 系统设置', level=2)
add_body('包含五个功能Tab：')

add_image_placeholder('系统设置界面截图（5个Tab）')
add_caption('图 8 系统设置界面')

add_table(
    ['Tab', '功能'],
    [
        ['用户权限', '用户增删改查、角色分配、启用/禁用'],
        ['角色管理', '管理员/试验工程师/只读用户三种内置角色'],
        ['操作日志', '按类型/时间筛选、清理、导出、保留天数配置'],
        ['数据备份', '手动备份、还原、自动备份间隔与保留策略、历史查看'],
        ['数据库高可用', '主从连接状态、故障切换参数、SQL Server连接配置'],
    ]
)

# === 10. 数据库状态与高可用 ===
doc.add_heading('10. 数据库状态与高可用', level=2)
add_body('左侧导航栏底部实时显示主库/从库连接状态（绿=在线，灰=离线），标记"当前"的为正在使用的库。')
add_body('主库故障时自动切换至从库，恢复后自动切回。切换期间通过磁盘缓冲机制确保数据零丢失。')
add_body('支持可配置间隔的自动备份和保留天数策略，备份状态在首页实时显示。')

# === 11. 附录 ===
doc.add_heading('11. 附录：操作提示', level=2)

doc.add_heading('11.1 完整试验流程', level=3)
add_body('1) 试验对象 → 确认项目/机组/对象/装置已登记')
add_body('2) 试验路径 → 配置试验配方')
add_body('3) 试验对象管理 → 为叶子节点关联默认试验路径')
add_body('4) 实时监视 → 选择对象 → 连接PLC → 开始监视')
add_body('5) 等待完成 → 停止监视（自动保存并判定）')
add_body('6) 试验记录 → 查看曲线回放 / 导出报告')

doc.add_heading('11.2 注意事项', level=3)
add_list('停止监视前不要关闭软件（虽有自动保存，建议正常停止）')
add_list('定期检查左下角数据库状态，确保主库在线')
add_list('变量配置修改后必须点"保存配置"，否则重启后丢失')
add_list('测量装置必须先在台账中登记并启用')
add_list('Y轴按可见窗口自适应，历史尖峰滑出后自动缩小')
add_list('配方修改不影响历史记录（每条记录保存导入时的配方快照）')
add_list('删除操作不可恢复，批量删除请谨慎')
add_list('远程数据库连接需确保TCP/IP启用、防火墙放行1433端口')

# ============ 保存 ============
OUTPUT = 'doc/用户操作手册.docx'
doc.save(OUTPUT)
print(f'已保存到：{OUTPUT}')
print(f'文件大小：{os.path.getsize(OUTPUT)/1024:.1f} KB')
