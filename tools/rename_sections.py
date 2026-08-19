# -*- coding: utf-8 -*-
import sys, os, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document('doc/数据分析-管理软件设计方案-A(2).docx')
count = 0

for i, p in enumerate(doc.paragraphs):
    t = p.text

    # 1. 模块列表项：基础台账（试验对象）→ 试验对象
    if '基础台账（试验对象）' in t:
        new_text = t.replace('基础台账（试验对象）', '试验对象')
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = new_text
            count += 1
            print(f'  [{i}] 模块名修正: {new_text[:60]}')

    # 2. 标题/正文：基础台账 → 试验对象（排除已处理的"基础台账（试验对象）"）
    elif '基础台账' in t and '基础台账（试验对象）' not in t:
        new_text = t.replace('基础台账', '试验对象')
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = new_text
            count += 1
            print(f'  [{i}] 基础台账→试验对象: {new_text[:60]}')

# 3. 同时修正图2标题（可能已被替换过）
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # 确保图片占位文字一致
    if '试验对象管理界面截图' in t and '图片占位' in t:
        pass  # 已正确
    elif '基础台账管理界面截图' in t:
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = t.replace('基础台账管理界面截图', '试验对象管理界面截图（4个Tab）')
            count += 1

# 4. 模块列表项里 "系统设置" 补充数据库高可用
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '系统设置: 包括用户管理' in t and '数据库高可用' not in t:
        for r in p.runs: r.text = ''
        if p.runs:
            p.runs[0].text = '系统设置: 包括用户管理、角色权限、备份策略、日志审计、数据库高可用等。'
            count += 1
            print(f'  [{i}] 系统设置描述补充')

# 保存到临时文件再复制
tmp = 'doc/数据分析-管理软件设计方案-A(2)_tmp.docx'
doc.save(tmp)
shutil.move(tmp, 'doc/数据分析-管理软件设计方案-A(2).docx')
print(f'\n修改了 {count} 处，已保存')

# 验证
doc2 = Document('doc/数据分析-管理软件设计方案-A(2).docx')
all_t = '\n'.join(p.text for p in doc2.paragraphs)

# 检查没有"基础台账"残留
if '基础台账' in all_t:
    print('❌ 仍有"基础台账"残留')
    for p in doc2.paragraphs:
        if '基础台账' in p.text:
            print(f'  {p.text[:80]}')
else:
    print('✅ "基础台账"已全部替换为"试验对象"')

# 检查导航对应
nav_items = ['首页概览', '试验对象', '试验路径', '试验记录', '实时监视', '数据分析', '系统设置']
print('\n导航项覆盖检查：')
for nav in nav_items:
    found = nav in all_t
    print(f'  {"✅" if found else "❌"} {nav}')
