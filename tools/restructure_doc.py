# -*- coding: utf-8 -*-
import sys, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

doc = Document('doc/数据分析-管理软件设计方案-A(2).docx')

# ============ 工具函数 ============
def make_para(text, template_para):
    new_p = deepcopy(template_para._element)
    runs = new_p.findall(qn('w:r'))
    first = True
    for r in runs:
        for t in r.findall(qn('w:t')):
            if first:
                t.text = text
                first = False
            else:
                t.text = ''
    return new_p

def find_heading3(text_contains):
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 3' and text_contains in p.text:
            return i
    return None

def find_heading1(text_contains):
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and text_contains in p.text:
            return i
    return None

def remove_paragraph(p):
    p._element.getparent().remove(p._element)

def get_next_non_empty_para(start_idx):
    for j in range(start_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].text.strip():
            return j
    return None

# ============ 获取模板 ============
heading3_tmpl = None
normal_tmpl = None
caption_tmpl = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 3' and not heading3_tmpl:
        heading3_tmpl = p
    if p.style.name == 'Normal' and p.runs and not normal_tmpl:
        normal_tmpl = p
    if p.style.name == 'Caption' and not caption_tmpl:
        caption_tmpl = p

# ============ 1. 删除"数据库高可用设计"整章 ============
db_start = find_heading1('数据库高可用设计')
if db_start is not None:
    # 找到下一个 Heading 1
    db_end = None
    for j in range(db_start + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 1':
            db_end = j
            break
    if db_end is None:
        db_end = len(doc.paragraphs)

    # 删除从 db_start 到 db_end 之前的所有段落
    removed = 0
    for i in range(db_end - 1, db_start - 1, -1):
        remove_paragraph(doc.paragraphs[i])
        removed += 1
    print(f'1. 删除"数据库高可用设计"章节：{removed} 段')
else:
    print('1. "数据库高可用设计"章节未找到，跳过')

# ============ 2. 精简非核心页面 ============

# --- 试验路径管理页面：精简为2段 ---
path_idx = find_heading3('试验路径管理页面')
if path_idx:
    # 收集从标题到下一个 Heading 3 之间的段落
    end_idx = None
    for j in range(path_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 3':
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    # 删除所有正文（保留标题）
    for i in range(end_idx - 1, path_idx, -1):
        p = doc.paragraphs[i]
        if p.style.name != 'Heading 3' and p.style.name != 'Caption':
            remove_paragraph(p)
        elif p.style.name == 'Caption':
            remove_paragraph(p)

    # 在标题后插入精简内容
    body = doc.paragraphs[path_idx]._element.getparent()
    title_elem = doc.paragraphs[path_idx]._element
    idx = list(body).index(title_elem) + 1

    brief = [
        '试验路径管理页面用于管理试验配方，定义泄漏率限值、预充压压力、阀门规格等试验参数。页面提供搜索、启用过滤、CSV导入导出、以及配方的增删改功能。配方列表以表格形式展示，包含名称、序号、系统、贯穿件直径、阀门编号、泄漏率限值、预充压P2、启用状态等字段。',
        '配方编辑对话框分三个区域：基础信息（名称、系统、启用状态、备注）、阀门参数（贯穿件直径、阀门编号、公称直径）、试验参数（泄漏率设计最大值、预充压P2）。每次编辑自动创建版本快照（RecipeVersion），试验记录中保存导入时的配方快照（JSON），后续修改配方不影响历史记录。',
        '【图片占位：试验路径管理界面截图】',
    ]
    for text in reversed(brief):
        body.insert(idx, make_para(text, normal_tmpl))
    print('2a. 精简试验路径管理页面 → 2段+图占位')

# --- 数据分析界面：精简为1段 ---
da_idx = find_heading3('数据分析界面')
if da_idx:
    end_idx = None
    for j in range(da_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 3':
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    for i in range(end_idx - 1, da_idx, -1):
        p = doc.paragraphs[i]
        if p.style.name != 'Heading 3':
            remove_paragraph(p)

    body = doc.paragraphs[da_idx]._element.getparent()
    title_elem = doc.paragraphs[da_idx]._element
    idx = list(body).index(title_elem) + 1

    brief = [
        '数据分析界面提供五个分析维度：故障趋势（按阀门类型统计合格/不合格数，堆叠柱状图展示）、合格率统计（各阀门合格率，≥95%合格/80%-95%注意/<80%不合格）、泄漏率趋势（多系列曲线对比）、阀门试验次数（Top 50排名+Top 20柱状图）、机组合格情况（各机组合格率对比）。支持按项目/机组/系统/时间范围筛选，可导出多Sheet的Excel文件。',
        '【图片占位：数据分析界面截图（5个Tab）】',
    ]
    for text in reversed(brief):
        body.insert(idx, make_para(text, normal_tmpl))
    print('2b. 精简数据分析界面 → 1段+图占位')

# --- 系统管理界面：精简为1段 ---
sm_idx = find_heading3('系统管理界面')
if sm_idx:
    end_idx = None
    for j in range(sm_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 3' or doc.paragraphs[j].style.name == 'Heading 1':
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    for i in range(end_idx - 1, sm_idx, -1):
        p = doc.paragraphs[i]
        if p.style.name != 'Heading 3' and p.style.name != 'Heading 1':
            remove_paragraph(p)

    body = doc.paragraphs[sm_idx]._element.getparent()
    title_elem = doc.paragraphs[sm_idx]._element
    idx = list(body).index(title_elem) + 1

    brief = [
        '系统管理界面包含五个功能Tab：用户权限（增删改查、角色分配、启用/禁用）、角色管理（管理员/试验工程师/只读用户三种内置角色）、操作日志（按类型/时间筛选、日志清理与导出、保留天数配置）、数据备份（手动备份、数据库还原、自动备份间隔与保留策略、备份历史查看）、数据库高可用配置（主从连接状态、故障切换参数、SQL Server连接配置）。',
        '【图片占位：系统管理界面截图（5个Tab）】',
    ]
    for text in reversed(brief):
        body.insert(idx, make_para(text, normal_tmpl))
    print('2c. 精简系统管理界面 → 1段+图占位')

# --- 首页概览：精简为1段 ---
ov_idx = find_heading3('软件首页概览')
if ov_idx:
    end_idx = None
    for j in range(ov_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 3':
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    for i in range(end_idx - 1, ov_idx, -1):
        p = doc.paragraphs[i]
        if p.style.name != 'Heading 3':
            remove_paragraph(p)

    body = doc.paragraphs[ov_idx]._element.getparent()
    title_elem = doc.paragraphs[ov_idx]._element
    idx = list(body).index(title_elem) + 1

    brief = [
        '首页概览采用"左侧导航 + 顶部核心指标 + 中部业务监控 + 底部台账概况"布局，以6个KPI卡片（试验对象数、测量装置数、历史记录数、本月合格率、待处理异常数、最近备份时间）、试验记录预览表、最近导入详情、9项台账统计指标和系统维护状态（数据库连接/备份/同步状态）为主体，满足全局数据一屏总览的需求。',
        '【图片占位：首页概览界面截图】',
    ]
    for text in reversed(brief):
        body.insert(idx, make_para(text, normal_tmpl))
    print('2d. 精简首页概览 → 1段+图占位')

# ============ 3. 展开核心页面 ============

# --- 试验对象管理页面（重点展开）---
to_idx = find_heading3('试验对象管理页面')
if to_idx:
    end_idx = None
    for j in range(to_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 3':
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    for i in range(end_idx - 1, to_idx, -1):
        p = doc.paragraphs[i]
        if p.style.name != 'Heading 3':
            remove_paragraph(p)

    body = doc.paragraphs[to_idx]._element.getparent()
    title_elem = doc.paragraphs[to_idx]._element
    idx = list(body).index(title_elem) + 1

    detailed = [
        '试验对象管理页面是系统的基础数据管理核心，采用四Tab布局，分别覆盖项目/机组管理、试验对象路径树管理、测量装置台账管理和试验报告导出功能。所有基础数据在同一页面内快速切换维护，实现试验数据的标准化、结构化管理。',

        'Tab 1 - 项目/机组管理：左侧为项目列表，右侧为机组列表（按选中项目过滤）。支持项目和机组的新增、编辑、删除操作，项目编码自动生成（格式P{年月}{序号}）。支持通过"批量导入数据"按钮从CSV文件夹批量导入项目和机组信息，进度条实时显示导入进度。',

        'Tab 2 - 试验对象路径树管理：采用四级层级树结构——系统(System) → 贯穿件(Penetration) → 阀门(Valve)/其他部件(OtherComponent)。页面顶部提供项目/机组下拉框范围过滤和关键字搜索定位功能。左侧为可展开折叠的试验对象树，底部提供四个快速新建按钮（新建系统/贯穿件/阀门/其他部件）。右侧显示选中节点的详细信息（编号、名称、类型、泄漏率限值、试验压力、父节点、备注）及操作按钮（修改、导入数据、导出数据、删除）。下方展示该对象的试验统计（累计试验次数、合格/不合格次数、最近结果）和关联的试验路径配方信息。叶子节点可配置默认关联试验路径，用于后续数据导入时自动匹配配方。',

        'Tab 3 - 测量装置台账管理：以数据表格形式展示所有测量装置信息，包括装置编号、装置名称、IP、序列号、主通信方式（USB/RJ45/RS232/RS485）、启用状态、最近连接状态、最近同步时间、最近导入时间。支持按通信方式和启用状态筛选，支持关键字搜索。提供装置的新增、编辑、删除操作。新增装置时编号自动生成（格式DEV-{时间戳}）。只有启用状态的装置才能在实时监视中被选择。',

        'Tab 4 - 报告导出：支持导出全部试验记录、本月试验记录、本月合格记录、本月不合格记录四种范围，支持Excel和PDF两种格式。可自定义文件名和导出目录。提供"导出Excel"和"导出PDF"快速导出按钮，一键生成试验报告。Excel报告包含试验记录汇总表和统计数据，PDF报告每条记录独立一页。',

        '【图片占位：试验对象管理界面截图（4个Tab）】',
    ]
    for text in reversed(detailed):
        body.insert(idx, make_para(text, normal_tmpl))
    print('3a. 展开试验对象管理页面 → 5段+图占位')

# --- 试验记录界面（重点展开）---
tr_idx = find_heading3('试验记录界面')
if tr_idx:
    end_idx = None
    for j in range(tr_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 3' or doc.paragraphs[j].style.name == 'Heading 1':
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    for i in range(end_idx - 1, tr_idx, -1):
        p = doc.paragraphs[i]
        if p.style.name != 'Heading 3' and p.style.name != 'Heading 1':
            remove_paragraph(p)

    body = doc.paragraphs[tr_idx]._element.getparent()
    title_elem = doc.paragraphs[tr_idx]._element
    idx = list(body).index(title_elem) + 1

    detailed = [
        '试验记录界面是数据管理软件的核心业务模块，用于对试验全过程数据进行集中管理、查询、分析与追溯。系统通过试验记录模块实现试验全过程数字化存档，为数据分析、设备状态评估、异常追溯及技术归档提供完整支撑。',

        '查询区域位于页面顶部，支持按项目（级联过滤机组）、试验结果（全部/合格/不合格/未知）、时间范围、关键字（记录编号/试验对象/测量装置/数据包名称）等条件进行组合筛选。查询结果以分页数据表格形式展示，包含记录编号、项目、机组、对象编码、节点名称、最终泄漏率(Nml/min)、泄漏限值(Nml/min)、判定结果（合格绿色/不合格红色）、关联试验路径、测量装置、操作人员、试验时间、导入时间、备注等字段。支持表头全选复选框批量选择记录。',

        '过程曲线回放功能位于页面底部。选中一条记录后，自动加载该记录的三张趋势曲线图（压力MPa、温度℃、流量Nml/min），曲线按通道属性自动分组到对应图表。图表支持左键拖拽平移、滚轮缩放Y轴、鼠标悬停显示Tracker浮层（各通道当前值）。提供"从Xs到Ys（0=全部）"时间范围输入框，可裁剪显示指定时间段内的数据。右侧面板显示各通道图例（名称、最小-最大值范围、单位）及关联试验路径的完整参数信息（配方名称、系统、贯穿件直径、阀门编号、公称直径、泄漏率限值、预充压P2、备注）。',

        '支持批量操作：勾选多条记录后可"批量修改试验路径"（选择新配方，系统自动重新计算泄漏限值和合格/不合格判定）或"批量删除"（不可恢复，需二次确认）。双击单条记录可打开编辑对话框修改关联试验路径和备注。数据上传支持单文件导入（.json/.txt/.csv格式，自动识别试验对象并关联默认配方）和批量导入（选择文件夹自动解析匹配）。导出功能支持Excel、PDF、CSV格式，可一键生成标准化试验报告。',

        '【图片占位：试验记录界面截图（含过程曲线回放）】',
    ]
    for text in reversed(detailed):
        body.insert(idx, make_para(text, normal_tmpl))
    print('3b. 展开试验记录界面 → 4段+图占位')

# --- 实时监视界面（重点展开）---
rm_idx = find_heading3('实时监视界面')
if rm_idx:
    end_idx = None
    for j in range(rm_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].style.name == 'Heading 3' or doc.paragraphs[j].style.name == 'Heading 1':
            end_idx = j
            break
    if end_idx is None:
        end_idx = len(doc.paragraphs)

    for i in range(end_idx - 1, rm_idx, -1):
        p = doc.paragraphs[i]
        if p.style.name != 'Heading 3' and p.style.name != 'Heading 1':
            remove_paragraph(p)

    body = doc.paragraphs[rm_idx]._element.getparent()
    title_elem = doc.paragraphs[rm_idx]._element
    idx = list(body).index(title_elem) + 1

    detailed = [
        '实时监视界面是软件的核心功能模块，用于连接PLC实时采集压力、温度、流量等数据并以趋势曲线形式展示。页面分为三个区域：顶部控制区、中部变量表格、底部趋势曲线图。',

        '顶部控制区包含试验对象选择和PLC连接控制两部分。试验对象选择采用四级级联下拉框：项目 → 机组（按项目过滤） → 试验对象（按机组过滤，显示编码+名称） → 测量装置（仅显示启用状态的台账装置）。PLC连接区域提供IP地址输入框（默认127.0.0.1）、保存地址、连接PLC、断开PLC按钮，系统自动识别Modbus TCP和Siemens S7两种协议。连接成功后提供开始监视、停止监视、导出CSV按钮。连接状态以颜色指示（红色=未连接，绿色=已连接）。监视过程中连续3次读取失败自动触发重连，重连失败则停止监视并报告错误。',

        '中部实时变量表格支持在线编辑，列包括：显示开关（复选框控制曲线显隐）、颜色标识、变量名称、西门子地址（如DB15.DBD0）、寄存器地址（Modbus地址）、数据类型（double/int/float/real/ushort/dword）、单位（MPa/℃/Nml/min等）、最小值/最大值（显示范围）、当前值（只读）、更新时间（只读）、状态（正常/待连接/未读取到数据，只读）。修改后点击"保存配置"持久化到数据库。支持添加变量、删除变量操作。',

        '底部三张趋势曲线图按变量通道属性自动分组：压力图(MPa)分组含"pressure/压力"通道变量，温度图(℃)分组含"temp/温度"通道变量，流量图(Nml/min)分组含"flow/流量"通道变量及未归类通道。每张图表下方显示图例（通道名称、当前值、单位），支持勾选/取消复选框控制曲线显隐。图表交互：左键拖拽平移X轴、滚轮缩放Y轴、鼠标悬停显示Tracker浮层。通过"显示时长"输入框（默认600秒）和"自动"复选框控制视口——勾选自动时视口跟随最新数据滚动，Y轴按当前可见窗口自适应；取消自动时视口停在当前位置，可自由拖拽查看历史。所有已采集数据始终保留，不裁剪不删除。停止监视时系统自动计算最终泄漏率（取M1/M2所有采样点的最大值），并根据泄漏限值判定合格/不合格。',

        '数据安全机制：即使不主动停止监视，系统也会周期自动保存已采集数据（间隔随数据量动态调整：10s → 30s → 60s → 5min），防止意外关闭导致数据丢失。内存缓冲区上限约86400点（1秒间隔约24小时），超出后自动裁剪旧点。停止或关闭时同步保存最终版数据到数据库。',

        '【图片占位：实时监视界面截图（含趋势曲线和变量表格）】',
    ]
    for text in reversed(detailed):
        body.insert(idx, make_para(text, normal_tmpl))
    print('3c. 展开实时监视界面 → 5段+图占位')

# ============ 保存 ============
tmp = 'doc/数据分析-管理软件设计方案-A(2)_tmp.docx'
doc.save(tmp)
shutil.move(tmp, 'doc/数据分析-管理软件设计方案-A(2).docx')
print(f'\n已保存!')
